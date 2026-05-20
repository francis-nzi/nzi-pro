"""
API endpoints for job report generation.
Generates comprehensive PDF reports with emissions data.
"""

from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import HTMLResponse, Response
from pydantic import BaseModel, Field
import logging
import io
from datetime import datetime
import hashlib
import re
import os
import base64
import json
import urllib.parse
from pathlib import Path
from jinja2 import Environment, FileSystemLoader
from typing import Optional, Any

# Import heavyweight deps inside functions to avoid startup overhead and missing dep errors
# - matplotlib: Used for chart generation (in chart_generation module)
# - docraptor: Used for professional PDF generation
# - python-docx: Used for DOCX export
# - playwright: Used for Playwright-based PDF generation (legacy)

from core.database import get_conn
from api.auth import _current_user
from services.monthly_emissions import JobMonthlyEmissionsResolver
from services.emissions_reporting import combined_row_metrics
from api.report_template_routes import (
    _get_job_report_metadata,
    _build_report_render_values,
    _ensure_report_template_schema,
)
from api.chart_generation import (
    generate_report_assets,
    validate_report_assets,
    REQUIRED_ASSETS,
)
from api.job_files_routes import (
    UPLOAD_DIR as JOB_FILES_UPLOAD_DIR,
    _drive_base_path as _files_drive_base_path,
    _ensure_job_files_table,
    _graph_request as _files_graph_request,
    _graph_download as _files_graph_download,
    _graph_raw_request as _files_graph_raw_request,
    _graph_token as _files_graph_token,
    _joined_remote_path as _files_joined_remote_path,
    _onedrive_ensure_folder as _files_onedrive_ensure_folder,
    _onedrive_enabled as _files_onedrive_enabled,
)
from services.playwright_browser import ensure_playwright_browser
from services.report_actions import get_job_report_actions_payload
from services.report_drafting import generate_report_section_draft
from services.client_benchmark import ensure_client_benchmark_columns, get_client_benchmark_metrics
from services.audit_log import record_audit_event
from api.job_data_output_routes import _load_data_output_rows, _build_scope_summary

logger = logging.getLogger(__name__)

# DocRaptor configuration
DOCRAPTOR_API_KEY = os.getenv('DOCRAPTOR_API_KEY', 'YOUR_TEST_API_KEY_GENERATES_WATERMARKS')
# Production mode when a real key is configured; test mode (watermarked) otherwise
_DOCRAPTOR_TEST_MODE = (DOCRAPTOR_API_KEY == 'YOUR_TEST_API_KEY_GENERATES_WATERMARKS')

ACTIVITY_GROUP_ORDER = [
    'Energy',
    'Business Travel',
    'Employee Commuting',
    'Purchased Goods & Services (PG&S)',
    'Other Emissions',
]

ACTIVITY_GROUP_COLORS = {
    'Energy': '#4472C4',
    'Business Travel': '#ED7D31',
    'Employee Commuting': '#A5A5A5',
    'Purchased Goods & Services (PG&S)': '#8064A2',
    'Other Emissions': '#70AD47',
}


def _table_columns(con, table_name: str) -> set[str]:
    try:
        rows = con.execute(
            """
            SELECT column_name
            FROM information_schema.columns
            WHERE table_schema = CURRENT_SCHEMA()
              AND table_name = %s
            """,
            [table_name],
        ).fetchall()
    except Exception:
        logger.debug("Failed to read table columns for report helper; returning empty set", exc_info=True)
        return set()
    return {str(row[0]).strip().lower() for row in rows if row and row[0] is not None}


def _factor_category_expr(con) -> str:
    return "fl.category" if "category" in _table_columns(con, "factor_lookup") else "NULL::text"

# Guards so DDL migrations run only once per process lifetime, not on every request.
_glossary_schema_seeded: bool = False
_report_versions_schema_seeded: bool = False
_report_drafts_schema_seeded: bool = False

DEFAULT_GLOSSARY_ENTRIES: list[dict[str, str]] = [
    {"term": "Absolute Emissions", "definition": "The total amount of greenhouse gasses calculated, measured in tonnes of CO2e."},
    {"term": "Benchmark Data", "definition": "The chosen 12-month period that sets the calculated emissions that need to be mitigated and/or offset."},
    {"term": "Carbon Reduction", "definition": "Reduction in measured CO2e emissions."},
    {"term": "Carbon Emissions (Gross)", "definition": "CO2e emissions from Company activities."},
    {"term": "Carbon Reduction Plan", "definition": "Plan to reduce CO2e emissions over a period of time, updated annually."},
    {"term": "tCO₂e", "definition": "Tonnes of carbon dioxide equivalent - a standard unit for measuring carbon footprints."},
    {"term": "Scope 1", "definition": "Direct GHG emissions from sources owned or controlled by the organisation."},
    {"term": "Scope 2", "definition": "Indirect GHG emissions from purchased electricity, steam, heating and cooling."},
    {"term": "Scope 3", "definition": "All other indirect emissions in the value chain."},
    {"term": "GHG Protocol", "definition": "The globally recognized framework for measuring and managing greenhouse gas emissions."},
    {"term": "Net Zero", "definition": "Achieving a balance between greenhouse gas emissions produced and emissions removed from the atmosphere."},
    {"term": "Baseline Year", "definition": "The historical reference year against which emissions reductions are measured."},
    {"term": "SECR", "definition": "Streamlined Energy and Carbon Reporting - UK reporting requirement"},
    {"term": "Intensity Metric", "definition": "Emissions normalised by business activity (e.g., per employee, per GBP revenue)."},
    {"term": "Emission Factor", "definition": "A coefficient that converts activity data into GHG emissions."},
]

router = APIRouter()


class GenerateReportPayload(BaseModel):
    template_id: Optional[int] = None
    version_id: Optional[int] = None
    save_version: bool = False
    report_version_status: Optional[str] = None
    report_version_label: Optional[str] = None
    report_version_notes: Optional[str] = None


class ReportVersionUpdatePayload(BaseModel):
    status: Optional[str] = None
    version_label: Optional[str] = None
    notes: Optional[str] = None


class GenerateReportDraftPayload(BaseModel):
    section_key: str
    provider: Optional[str] = "anthropic"
    model: Optional[str] = None
    template_key: Optional[str] = None
    sibling_drafts: Optional[dict] = None


class ReportDraftSectionPayload(BaseModel):
    section_key: str
    section_title: Optional[str] = None
    draft_text: Optional[str] = None
    draft_json: Optional[dict[str, Any]] = None
    provider: Optional[str] = None
    model: Optional[str] = None
    confidence: Optional[str] = None
    origin: Optional[str] = None


class SaveReportDraftsPayload(BaseModel):
    template_key: Optional[str] = None
    sections: list[ReportDraftSectionPayload] = Field(default_factory=list)


_SYSTEM_LOGO_PATH = (
    Path(__file__).resolve().parents[1] / "frontend" / "public" / "uploads" / "system" / "nzi-logo.png"
)


def _get_nzi_logo_src() -> str:
    """Return NZI logo as data URI. Checks DB first (redeploy-safe), falls back to filesystem."""
    # 1. Try database (logo uploaded via Admin â†’ System Settings)
    try:
        with get_conn() as con:
            rows = con.execute(
                "SELECT setting_key, setting_value FROM system_settings WHERE setting_key IN (%s, %s)",
                ["nzi_logo_b64", "nzi_logo_mime"],
            ).fetchall()
            values = {str(r[0] or ""): r[1] for r in rows}
            logo_b64 = values.get("nzi_logo_b64") or ""
            mime = str(values.get("nzi_logo_mime") or "image/png").strip()
            if logo_b64:
                return f"data:{mime};base64,{logo_b64}"
    except Exception:
        logger.debug("Falling back to filesystem logo lookup after DB lookup failed", exc_info=True)
    # 2. Try local file (dev / first-run before any upload)
    try:
        if _SYSTEM_LOGO_PATH.exists():
            return "data:image/png;base64," + base64.b64encode(_SYSTEM_LOGO_PATH.read_bytes()).decode("ascii")
    except Exception:
        logger.debug("Falling back to empty logo after filesystem logo lookup failed", exc_info=True)
    return ""


def _get_client_logo_src(logo_url: str | None) -> str:
    """Return client logo as data URI or absolute URL. Returns empty string if unavailable."""
    if not logo_url:
        return ""
    logo_url = str(logo_url).strip()
    # External URL â€” use directly
    if logo_url.startswith("http://") or logo_url.startswith("https://"):
        return logo_url
    # Relative path to uploaded file â€” try to read and base64-encode
    try:
        from pathlib import Path as _Path
        project_root = _Path(__file__).resolve().parents[1]
        # Strip leading slash for path join
        rel = logo_url.lstrip("/")
        file_path = project_root / "frontend" / "public" / rel
        if file_path.exists():
            suffix = file_path.suffix.lower()
            mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                        ".gif": "image/gif", ".svg": "image/svg+xml", ".webp": "image/webp"}
            mime = mime_map.get(suffix, "image/png")
            return f"data:{mime};base64," + base64.b64encode(file_path.read_bytes()).decode("ascii")
    except Exception:
        logger.debug("Falling back to original client logo URL after local file lookup failed", exc_info=True)
    return logo_url


def _normalize_render_value(value):
    """Normalize DB values for template rendering (handle NULL/NaN)."""
    if value is None:
        return ""
    try:
        # NaN check (NaN != NaN)
        if isinstance(value, float) and value != value:
            return ""
    except Exception:
        logger.debug("Unable to normalize render value; returning original value", exc_info=True)
    return value


def _resolve_builtin_template_name(selection: dict | None) -> str:
    """Choose a bundled template file when no custom template_content exists."""
    if not selection:
        return "report_template.html"

    hint = (
        f"{selection.get('template_type', '')} "
        f"{selection.get('report_type', '')} "
        f"{selection.get('template_key', '')} "
        f"{selection.get('template_name', '')}"
    ).lower()

    # CRP templates use the previously developed interactive layout with full section set.
    if (
        ("crp" in hint)
        or ("carbon_reduction_plan" in hint)
        or ("carbon reduction plan" in hint)
        or ("ppn_006" in hint)
        or ("ppn 006" in hint)
        or ("ppn06/21" in hint)
    ):
        return "interactive_report.html"

    # Only explicitly professional templates should use the professional layout.
    if "professional" in hint:
        return "professional_report.html"

    return "report_template.html"


def _use_forced_builtin_template(selection: dict | None) -> bool:
    """Force bundled template only for explicitly professional templates."""
    if not selection:
        return False
    hint = (
        f"{selection.get('template_type', '')} "
        f"{selection.get('report_type', '')} "
        f"{selection.get('template_key', '')} "
        f"{selection.get('template_name', '')}"
    ).lower()
    return (
        ("professional" in hint)
        or ("crp" in hint)
        or ("carbon_reduction_plan" in hint)
        or ("carbon reduction plan" in hint)
        or ("ppn_006" in hint)
        or ("ppn 006" in hint)
        or ("ppn06/21" in hint)
    )


def _resolve_template_for_render(env: Environment, selection: dict | None):
    """Resolve template source (built-in or DB content) for rendering."""
    if _use_forced_builtin_template(selection):
        return env.get_template(_resolve_builtin_template_name(selection))

    template_content = selection.get("template_content") if selection else None
    if template_content and str(template_content).strip():
        return env.from_string(str(template_content))

    template_name = _resolve_builtin_template_name(selection)
    return env.get_template(template_name)


def _get_template_selection(template_id: int, version_id: int | None = None):
    """Load template + version metadata, defaulting to latest version when missing."""
    with get_conn() as con:
        template_row = con.execute(
            """
            SELECT template_id, template_key, template_name, template_type, report_type, is_active
            FROM report_templates
            WHERE template_id = %s
            """,
            [int(template_id)],
        ).fetchone()

        if not template_row:
            return None
        if not bool(template_row[5]):
            return None

        if version_id is None:
            version_row = con.execute(
                """
                SELECT version_id, version_number, version_label, status, template_content
                FROM report_template_versions
                WHERE template_id = %s
                ORDER BY version_number DESC
                LIMIT 1
                """,
                [int(template_id)],
            ).fetchone()
        else:
            version_row = con.execute(
                """
                SELECT version_id, version_number, version_label, status, template_content
                FROM report_template_versions
                WHERE template_id = %s AND version_id = %s
                LIMIT 1
                """,
                [int(template_id), int(version_id)],
            ).fetchone()

    if not version_row:
        return {
            "template_id": int(template_row[0]),
            "template_key": template_row[1],
            "template_name": template_row[2],
            "template_type": template_row[3],
            "report_type": template_row[4],
            "version_id": None,
            "version_number": None,
            "version_label": None,
            "version_status": None,
            "template_content": None,
        }

    return {
        "template_id": int(template_row[0]),
        "template_key": template_row[1],
        "template_name": template_row[2],
        "template_type": template_row[3],
        "report_type": template_row[4],
        "version_id": int(version_row[0]),
        "version_number": int(version_row[1]) if version_row[1] is not None else None,
        "version_label": version_row[2],
        "version_status": version_row[3],
        "template_content": version_row[4],
    }


def _get_job_assigned_template_selection(job_id: int):
    """Get currently assigned template/version for a job."""
    with get_conn() as con:
        row = con.execute(
            """
            SELECT a.template_id, a.version_id
            FROM job_report_template_assignments a
            WHERE a.job_id = %s
            """,
            [int(job_id)],
        ).fetchone()

    if not row:
        return None
    if row[0] is None:
        return None

    return _get_template_selection(int(row[0]), int(row[1]) if row[1] is not None else None)


def _ensure_report_versions_schema(con) -> None:
    global _report_versions_schema_seeded
    if _report_versions_schema_seeded:
        return
    con.execute(
        """
        CREATE TABLE IF NOT EXISTS job_report_versions (
          report_version_id BIGINT GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
          org_id TEXT,
          job_id INTEGER NOT NULL REFERENCES jobs(job_id) ON DELETE CASCADE,
          client_db_id INTEGER NOT NULL REFERENCES clients(db_id) ON DELETE CASCADE,
          version_number INTEGER NOT NULL,
          version_label VARCHAR,
          status VARCHAR NOT NULL DEFAULT 'review',
          report_format VARCHAR NOT NULL DEFAULT 'pdf',
          template_id INTEGER REFERENCES report_templates(template_id) ON DELETE SET NULL,
          version_id INTEGER REFERENCES report_template_versions(version_id) ON DELETE SET NULL,
          file_id INTEGER REFERENCES job_files(file_id) ON DELETE SET NULL,
          file_name TEXT,
          file_path TEXT,
          storage_provider VARCHAR DEFAULT 'local',
          external_item_id VARCHAR,
          external_web_url TEXT,
          external_path TEXT,
          data_hash VARCHAR,
          snapshot_json TEXT,
          notes TEXT,
          generated_at TIMESTAMP DEFAULT NOW(),
          generated_by VARCHAR,
          reviewed_at TIMESTAMP,
          reviewed_by VARCHAR,
          finalized_at TIMESTAMP,
          finalized_by VARCHAR,
          superseded_at TIMESTAMP,
          superseded_by VARCHAR,
          UNIQUE(job_id, version_number)
        )
        """
    )
    _ddl_alters = [
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS org_id TEXT",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS version_label VARCHAR",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS file_name TEXT",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS file_path TEXT",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS storage_provider VARCHAR DEFAULT 'local'",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS external_item_id VARCHAR",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS external_web_url TEXT",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS external_path TEXT",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS data_hash VARCHAR",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS snapshot_json TEXT",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS notes TEXT",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS generated_by VARCHAR",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS reviewed_at TIMESTAMP",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS reviewed_by VARCHAR",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS finalized_at TIMESTAMP",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS finalized_by VARCHAR",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS superseded_at TIMESTAMP",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS superseded_by VARCHAR",
        "ALTER TABLE job_report_versions ADD COLUMN IF NOT EXISTS report_format VARCHAR DEFAULT 'pdf'",
        "CREATE INDEX IF NOT EXISTS job_report_versions_job_idx ON job_report_versions (job_id, version_number DESC)",
        "CREATE INDEX IF NOT EXISTS job_report_versions_client_idx ON job_report_versions (client_db_id, generated_at DESC)",
        "CREATE INDEX IF NOT EXISTS job_report_versions_org_idx ON job_report_versions (org_id, job_id, version_number DESC)",
    ]
    for _stmt in _ddl_alters:
        try:
            con.execute(_stmt)
        except Exception:
            logger.debug("Ignoring job_report_versions DDL compatibility failure", exc_info=True)
    try:
        con.execute(
            """
            UPDATE job_report_versions jrv
            SET org_id = COALESCE(jrv.org_id, c.org_id)
            FROM jobs j
            JOIN clients c ON c.db_id = j.client_db_id
            WHERE j.job_id = jrv.job_id
              AND jrv.org_id IS NULL
            """
        )
    except Exception:
        logger.debug("Ignoring job_report_versions org backfill failure", exc_info=True)
    _report_versions_schema_seeded = True


def _ensure_report_drafts_schema(con) -> None:
    global _report_drafts_schema_seeded
    if _report_drafts_schema_seeded:
        return
    con.execute(
        """
        CREATE TABLE IF NOT EXISTS job_report_drafts (
          draft_id INTEGER GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
          org_id TEXT,
          job_id INTEGER NOT NULL REFERENCES jobs(job_id) ON DELETE CASCADE,
          template_id INTEGER NOT NULL REFERENCES report_templates(template_id) ON DELETE CASCADE,
          version_id INTEGER NOT NULL REFERENCES report_template_versions(version_id) ON DELETE CASCADE,
          section_key VARCHAR NOT NULL,
          section_title VARCHAR,
          draft_text TEXT NOT NULL DEFAULT '',
          draft_json TEXT,
          evidence_hash VARCHAR,
          provider VARCHAR,
          model VARCHAR,
          confidence VARCHAR,
          status VARCHAR DEFAULT 'draft',
          revision_number INTEGER NOT NULL DEFAULT 1,
          created_at TIMESTAMP DEFAULT NOW(),
          updated_at TIMESTAMP DEFAULT NOW(),
          created_by VARCHAR,
          updated_by VARCHAR,
          UNIQUE(job_id, template_id, version_id, section_key)
        )
        """
    )
    con.execute("ALTER TABLE job_report_drafts ADD COLUMN IF NOT EXISTS org_id TEXT")
    con.execute(
        """
        CREATE INDEX IF NOT EXISTS job_report_drafts_lookup_idx
        ON job_report_drafts (job_id, template_id, version_id, section_key, updated_at DESC)
        """
    )
    con.execute(
        """
        CREATE INDEX IF NOT EXISTS job_report_drafts_org_idx
        ON job_report_drafts (org_id, job_id, template_id, version_id, section_key, updated_at DESC)
        """
    )
    try:
        con.execute(
            """
            UPDATE job_report_drafts jrd
            SET org_id = COALESCE(jrd.org_id, c.org_id)
            FROM jobs j
            JOIN clients c ON c.db_id = j.client_db_id
            WHERE j.job_id = jrd.job_id
              AND jrd.org_id IS NULL
            """
        )
    except Exception:
        logger.debug("Ignoring job_report_drafts org backfill failure", exc_info=True)
    _report_drafts_schema_seeded = True


def _safe_path_segment(text: str | None) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._ -]+", "_", str(text or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    return cleaned or "item"


def _report_version_folder(job_data: dict[str, Any], version_number: int) -> str:
    client_db_id = int(job_data.get("client_db_id") or 0)
    client_name = _safe_path_segment(job_data.get("client_name") or "client").replace(" ", "_")
    job_number = _safe_path_segment(job_data.get("job_number") or f"job-{job_data.get('job_id') or 'unknown'}").replace(" ", "_")
    return f"client-{client_db_id}/{client_name}/job-{job_number}/report-versions/v{int(version_number)}"


def _upload_report_bytes_to_folder(*, folder: str, filename: str, content: bytes) -> dict[str, str | int | None]:
    token = _files_graph_token()
    drive_base = _files_drive_base_path(token)
    remote_folder = _files_joined_remote_path(folder)
    _files_onedrive_ensure_folder(token, remote_folder)

    full_path = f"{remote_folder}/{filename}" if remote_folder else f"/{filename}"
    encoded = Path(full_path).as_posix()
    encoded = re.sub(r"//+", "/", encoded)
    encoded = encoded.replace(" ", "%20")
    encoded = encoded.replace("#", "%23")
    encoded = encoded.replace("?", "%3F")

    if len(content) <= 4 * 1024 * 1024:
        meta = _files_graph_request(
            "PUT",
            f"{drive_base}/root:{encoded}:/content",
            token,
            body=content,
            content_type="application/octet-stream",
        )
    else:
        session = _files_graph_request(
            "POST",
            f"{drive_base}/root:{encoded}:/createUploadSession",
            token,
            body=json.dumps(
                {
                    "item": {
                        "@microsoft.graph.conflictBehavior": "replace",
                        "name": filename,
                    }
                }
            ).encode("utf-8"),
            content_type="application/json",
        )
        upload_url = str(session.get("uploadUrl") or "").strip()
        if not upload_url:
            raise HTTPException(status_code=500, detail="Failed to create OneDrive upload session for report artifact")

        chunk_size = 5 * 1024 * 1024
        start = 0
        meta: dict[str, object] = {}
        while start < len(content):
            end = min(start + chunk_size, len(content)) - 1
            chunk = content[start : end + 1]
            raw = _files_graph_raw_request(
                "PUT",
                upload_url,
                body=chunk,
                headers={
                    "Content-Length": str(len(chunk)),
                    "Content-Range": f"bytes {start}-{end}/{len(content)}",
                    "Content-Type": "application/octet-stream",
                },
            )
            if raw:
                meta = json.loads(raw.decode("utf-8"))
            start = end + 1

    return {
        "storage_provider": "onedrive",
        "file_path": str(full_path),
        "external_item_id": str(meta.get("id") or ""),
        "external_web_url": str(meta.get("webUrl") or ""),
        "external_path": str(full_path),
        "file_size": len(content),
    }


def _store_report_version_artifact(
    *,
    con,
    job_data: dict[str, Any],
    version_number: int,
    pdf_bytes: bytes,
    generated_by: str,
    status: str,
    template_id: int | None,
    template_version_id: int | None,
    data_hash: str,
    snapshot_json: str,
    version_label: str | None = None,
    notes: str | None = None,
    report_format: str = "pdf",
) -> dict[str, Any]:
    _ensure_job_files_table(con)
    _ensure_report_versions_schema(con)

    client_db_id = int(job_data.get("client_db_id") or 0)
    org_id = str(job_data.get("org_id") or "").strip() or None
    safe_job_number = _safe_path_segment(job_data.get("job_number") or f"job-{job_data.get('job_id') or 'unknown'}").replace(" ", "_")
    safe_client_name = _safe_path_segment(job_data.get("client_name") or "client").replace(" ", "_")
    file_name = _safe_filename(
        f"{safe_job_number}-report-v{int(version_number)}",
        "pdf",
    )
    if _files_onedrive_enabled():
        storage_meta = _upload_report_bytes_to_folder(
            folder=_report_version_folder(job_data, version_number),
            filename=file_name,
            content=pdf_bytes,
        )
    else:
        folder_path = JOB_FILES_UPLOAD_DIR / f"client_{client_db_id}_{safe_client_name}" / f"job_{safe_job_number}" / "report-versions" / f"v{int(version_number)}"
        folder_path.mkdir(parents=True, exist_ok=True)
        file_path = folder_path / file_name
        file_path.write_bytes(pdf_bytes)
        storage_meta = {
            "storage_provider": "local",
            "file_path": str(file_path),
            "external_item_id": None,
            "external_web_url": None,
            "external_path": None,
            "file_size": len(pdf_bytes),
        }

    file_row = con.execute(
        """
        INSERT INTO job_files (
            job_id, row_id, file_type, file_name, file_path, file_size, mime_type,
            description, uploaded_by, storage_provider, external_item_id, external_web_url, external_path
        )
        VALUES (%s, NULL, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        RETURNING file_id
        """,
        [
            int(job_data["job_id"]),
            "generated_report",
            file_name,
            storage_meta.get("file_path"),
            storage_meta.get("file_size"),
            "application/pdf",
            f"Report version v{int(version_number)} ({status})",
            generated_by,
            storage_meta.get("storage_provider"),
            storage_meta.get("external_item_id"),
            storage_meta.get("external_web_url"),
            storage_meta.get("external_path"),
        ],
    ).fetchone()
    file_id = int(file_row[0]) if file_row else None

    version_label_value = str(version_label or f"v{int(version_number)}").strip() or f"v{int(version_number)}"
    version_row = con.execute(
        """
        INSERT INTO job_report_versions (
          org_id, job_id, client_db_id, version_number, version_label, status, report_format,
          template_id, version_id, file_id, file_name, file_path, storage_provider,
          external_item_id, external_web_url, external_path, data_hash, snapshot_json,
          notes, generated_by, reviewed_at, reviewed_by, finalized_at, finalized_by
        )
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                %s,
                CASE WHEN lower(%s) = 'review' THEN NOW() ELSE NULL END,
                CASE WHEN lower(%s) = 'review' THEN %s ELSE NULL END,
                CASE WHEN lower(%s) = 'final' THEN NOW() ELSE NULL END,
                CASE WHEN lower(%s) = 'final' THEN %s ELSE NULL END)
        RETURNING report_version_id
        """,
        [
            org_id,
            int(job_data["job_id"]),
            client_db_id,
            int(version_number),
            version_label_value,
            status,
            str(report_format or "pdf"),
            template_id,
            template_version_id,
            file_id,
            file_name,
            storage_meta.get("file_path"),
            storage_meta.get("storage_provider"),
            storage_meta.get("external_item_id"),
            storage_meta.get("external_web_url"),
            storage_meta.get("external_path"),
            data_hash,
            snapshot_json,
            notes,
            generated_by,
            status,
            status,
            generated_by,
            status,
            status,
            generated_by,
        ],
    ).fetchone()

    return {
        "report_version_id": int(version_row[0]) if version_row else None,
        "file_id": file_id,
        "file_name": file_name,
        "version_number": int(version_number),
        "version_label": version_label_value,
        "status": status,
        "storage_meta": storage_meta,
    }


def _normalize_report_version_status(status: str | None) -> str:
    norm = str(status or "review").strip().lower()
    return norm if norm in {"draft", "review", "final", "superseded", "archived"} else "review"


def _build_report_version_snapshot(
    *,
    job_data: dict[str, Any],
    selected_template: dict[str, Any] | None,
    template_variables: dict[str, Any],
    report_metadata: dict[str, Any],
    job_actions: dict[str, Any],
    scope_totals: dict[str, Any],
    benchmark_totals: dict[str, Any],
    categories: list[dict[str, Any]],
    activity_totals: dict[str, Any],
    site_breakdowns: dict[str, Any],
    benchmark_site_breakdowns: dict[str, Any],
    render_values: dict[str, Any],
    target_data: dict[str, Any],
    scope_comparison: dict[str, Any],
    activity_comparison: dict[str, Any],
    site_overall_comparison: dict[str, Any],
    assets: dict[str, Any],
    generation_date: str,
) -> tuple[dict[str, Any], str]:
    payload = {
        "job_data": job_data,
        "template": selected_template,
        "template_variables": template_variables,
        "report_metadata": report_metadata,
        "job_actions": job_actions,
        "scope_totals": scope_totals,
        "benchmark_totals": benchmark_totals,
        "categories": categories,
        "activity_totals": activity_totals,
        "site_breakdowns": site_breakdowns,
        "benchmark_site_breakdowns": benchmark_site_breakdowns,
        "render_values": render_values,
        "target_data": target_data,
        "scope_comparison": scope_comparison,
        "activity_comparison": activity_comparison,
        "site_overall_comparison": site_overall_comparison,
        "assets": assets,
        "asset_keys": sorted([str(key) for key in assets.keys()]),
        "generation_date": generation_date,
    }
    snapshot_json = json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str)
    data_hash = hashlib.sha256(snapshot_json.encode("utf-8")).hexdigest()
    return payload, data_hash


def _safe_int(v) -> int | None:
    """Convert v to int, returning None for NULL/NaN (pandas represents NULL integers as NaN)."""
    if v is None:
        return None
    try:
        f = float(v)
    except (TypeError, ValueError):
        return None
    import math
    if math.isnan(f) or math.isinf(f):
        return None
    return int(f)


def _safe_str(v) -> str | None:
    """Convert v to str, returning None for NULL/NaT."""
    if v is None:
        return None
    try:
        import pandas as pd
        if pd.isna(v):
            return None
    except (TypeError, ValueError):
        pass
    return str(v)


def _serialize_report_version_row(row: dict[str, Any]) -> dict[str, Any]:
    file_id = _safe_int(row.get("file_id"))
    job_id = _safe_int(row.get("job_id"))
    report_version_id = _safe_int(row.get("report_version_id"))
    return {
        "report_version_id": report_version_id,
        "job_id": job_id,
        "client_db_id": _safe_int(row.get("client_db_id")),
        "version_number": _safe_int(row.get("version_number")),
        "version_label": row.get("version_label"),
        "status": row.get("status"),
        "report_format": row.get("report_format"),
        "template_id": _safe_int(row.get("template_id")),
        "version_id": _safe_int(row.get("version_id")),
        "file_id": file_id,
        "file_name": row.get("file_name"),
        "file_path": row.get("file_path"),
        "storage_provider": row.get("storage_provider"),
        "external_item_id": row.get("external_item_id"),
        "external_web_url": row.get("external_web_url"),
        "external_path": row.get("external_path"),
        "data_hash": row.get("data_hash"),
        "notes": row.get("notes"),
        "generated_at": _safe_str(row.get("generated_at")),
        "generated_by": row.get("generated_by"),
        "reviewed_at": _safe_str(row.get("reviewed_at")),
        "reviewed_by": row.get("reviewed_by"),
        "finalized_at": _safe_str(row.get("finalized_at")),
        "finalized_by": row.get("finalized_by"),
        "superseded_at": _safe_str(row.get("superseded_at")),
        "superseded_by": row.get("superseded_by"),
        "download_url": f"/jobs/{job_id}/files/{file_id}/download" if job_id is not None and file_id is not None else None,
        "snapshot_url": f"/jobs/{job_id}/report-versions/{report_version_id}/snapshot-html"
        if job_id is not None and report_version_id is not None
        else None,
    }


def _load_report_version_snapshot(con, job_id: int, report_version_id: int) -> tuple[dict[str, Any], dict[str, Any]]:
    version_df = con.execute(
        """
        SELECT *
        FROM job_report_versions
        WHERE job_id = %s AND report_version_id = %s
        """,
        [int(job_id), int(report_version_id)],
    ).df()
    if version_df is None or version_df.empty:
        raise HTTPException(status_code=404, detail="Report version not found")

    version_row = version_df.iloc[0].to_dict()
    snapshot_raw = version_row.get("snapshot_json")
    if not snapshot_raw:
        raise HTTPException(status_code=404, detail="Report version snapshot is not available")
    try:
        snapshot_payload = json.loads(snapshot_raw)
    except Exception as exc:
        logger.warning("Failed to parse report version snapshot JSON", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to parse report version snapshot: {exc}")
    if not isinstance(snapshot_payload, dict):
        raise HTTPException(status_code=500, detail="Report version snapshot is invalid")
    return version_row, snapshot_payload


def _load_latest_final_report_version_snapshot(con, job_id: int) -> tuple[dict[str, Any], dict[str, Any]] | None:
    version_row = con.execute(
        """
        SELECT *
        FROM job_report_versions
        WHERE job_id = %s
          AND lower(COALESCE(status, '')) = 'final'
        ORDER BY COALESCE(finalized_at, generated_at) DESC NULLS LAST, version_number DESC, report_version_id DESC
        LIMIT 1
        """,
        [int(job_id)],
    ).fetchone()
    if not version_row:
        return None

    version_id = int(version_row[0]) if getattr(version_row, "__len__", None) else None
    if version_id is None:
        return None
    return _load_report_version_snapshot(con, int(job_id), version_id)


def _load_report_version_file_payload(con, job_id: int, file_id: int) -> dict[str, Any] | None:
    file_row = con.execute(
        """
        SELECT file_name, file_path, storage_provider, external_item_id, mime_type
        FROM job_files
        WHERE file_id = %s AND job_id = %s
        """,
        [int(file_id), int(job_id)],
    ).fetchone()
    if not file_row:
        return None

    file_name = str(file_row[0] or f"job-{job_id}-report.pdf")
    file_path = str(file_row[1] or "").strip()
    storage_provider = str(file_row[2] or "local").strip().lower()
    external_item_id = str(file_row[3] or "").strip()
    mime_type = str(file_row[4] or "application/octet-stream")

    if storage_provider == "onedrive":
        if not external_item_id:
            return None
        token = _files_graph_token()
        drive_base = _files_drive_base_path(token)
        content, content_type = _files_graph_download(
            f"{drive_base}/items/{urllib.parse.quote(external_item_id)}/content",
            token,
        )
        return {
            "file_name": file_name,
            "content": content,
            "media_type": content_type or mime_type or "application/octet-stream",
        }

    if not file_path or not os.path.exists(file_path):
        return None

    return {
        "file_name": file_name,
        "content": Path(file_path).read_bytes(),
        "media_type": mime_type or "application/octet-stream",
    }


def _render_report_snapshot_html(snapshot_payload: dict[str, Any]) -> str:
    template_meta = snapshot_payload.get("template") or snapshot_payload.get("render_template") or {}
    template_selection = None
    if isinstance(template_meta, dict):
        template_id = template_meta.get("template_id")
        version_id = template_meta.get("version_id")
        if template_id is not None and version_id is not None:
            try:
                template_selection = _get_template_selection(int(template_id), int(version_id))
            except Exception:
                logger.debug("Failed to resolve explicit report template selection; falling back to snapshot metadata", exc_info=True)
                template_selection = None
        elif template_id is not None:
            try:
                template_selection = _get_template_selection(int(template_id), None)
            except Exception:
                logger.debug("Failed to resolve report template selection without version; falling back to snapshot metadata", exc_info=True)
                template_selection = None
        if not template_selection:
            template_selection = template_meta
    template_dir = os.path.join(os.path.dirname(__file__), "templates")
    env = Environment(loader=FileSystemLoader(template_dir))
    template = _resolve_template_for_render(env, template_selection if isinstance(template_selection, dict) else None)

    categories = snapshot_payload.get("categories") or []
    activity_groups = snapshot_payload.get("activity_groups")
    activity_details = snapshot_payload.get("activity_details")
    if (not activity_groups or not activity_details) and categories:
        activity_groups, _, activity_details = _build_activity_grouping(categories)

    html_content = template.render(
        job_data=snapshot_payload.get("job_data") or {},
        scope_totals=snapshot_payload.get("scope_totals") or {},
        benchmark_totals=snapshot_payload.get("benchmark_totals") or {},
        categories=categories,
        activity_groups=activity_groups or {},
        activity_totals=snapshot_payload.get("activity_totals") or {},
        activity_details=activity_details or [],
        activity_group_order=ACTIVITY_GROUP_ORDER,
        activity_group_colors=ACTIVITY_GROUP_COLORS,
        assets=snapshot_payload.get("assets") or {},
        scope_chart_base64=(snapshot_payload.get("assets") or {}).get("scope_breakdown", ""),
        activity_chart_base64=(snapshot_payload.get("assets") or {}).get("activity_breakdown", ""),
        generation_date=snapshot_payload.get("generation_date") or "",
        template_variables=snapshot_payload.get("template_variables") or {},
        report_metadata=snapshot_payload.get("report_metadata") or {},
        job_actions=snapshot_payload.get("job_actions") or {},
        nzi_logo_src=_get_nzi_logo_src(),
        client_logo_src=_get_client_logo_src((snapshot_payload.get("job_data") or {}).get("logo_url")),
        render_values=snapshot_payload.get("render_values") or {},
        render_template=template_selection if isinstance(template_selection, dict) else {},
        glossary_terms=snapshot_payload.get("glossary_terms") or [],
        target_data=snapshot_payload.get("target_data") or {},
        site_breakdowns=snapshot_payload.get("site_breakdowns") or {},
        is_yoy_template=bool(snapshot_payload.get("is_yoy_template") or (template_meta.get("template_key") if isinstance(template_meta, dict) else "" ) == "crp_yoy_benchmark"),
        scope_comparison=snapshot_payload.get("scope_comparison") or {},
        benchmark_activity_totals=snapshot_payload.get("benchmark_activity_totals") or {},
        activity_comparison=snapshot_payload.get("activity_comparison") or {},
        benchmark_site_breakdowns=snapshot_payload.get("benchmark_site_breakdowns") or {},
        site_overall_comparison=snapshot_payload.get("site_overall_comparison") or {},
    )
    render_values = snapshot_payload.get("render_values") or {}
    return _apply_braced_placeholder_substitution(html_content, render_values)


def _render_html_to_pdf_bytes(html_content: str) -> bytes:
    # Render prefers DocRaptor in production so the web process doesn't have to
    # launch Chromium under memory pressure. Keep a local Playwright fallback for
    # developer machines where DocRaptor may not be configured.
    render_env = str(os.getenv("RENDER") or os.getenv("RENDER_SERVICE_ID") or "").strip()
    if render_env:
        import docraptor

        doc_api = docraptor.DocApi()
        doc_api.api_client.configuration.username = DOCRAPTOR_API_KEY
        return doc_api.create_doc(
            {
                "test": _DOCRAPTOR_TEST_MODE,
                "document_content": html_content,
                "document_type": "pdf",
                "name": "job-report.pdf",
                "prince_options": {
                    "media": "print",
                    "javascript": False,
                },
            }
        )

    from playwright.sync_api import sync_playwright

    ensure_playwright_browser()
    with sync_playwright() as p:
        browser = p.chromium.launch(args=["--no-sandbox", "--disable-dev-shm-usage", "--disable-gpu"])
        page = browser.new_page()
        page.set_content(html_content)
        pdf_bytes = page.pdf(
            format="A4",
            margin={
                "top": "20mm",
                "right": "20mm",
                "bottom": "20mm",
                "left": "20mm",
            },
            print_background=True,
        )
        browser.close()
    return pdf_bytes


def _get_job_assignment_template_and_version(job_id: int) -> tuple[int, int] | None:
    """Return the assigned template_id/version_id for a job, or None if incomplete."""
    with get_conn() as con:
        row = con.execute(
            """
            SELECT template_id, version_id
            FROM job_report_template_assignments
            WHERE job_id = %s
            """,
            [int(job_id)],
        ).fetchone()

    if not row:
        return None

    template_id = int(row[0]) if row[0] is not None else None
    version_id = int(row[1]) if row[1] is not None else None
    if template_id is None or version_id is None:
        return None

    return template_id, version_id


def _get_template_variable_values_for_render(job_id: int, template_id: int, version_id: int | None = None):
    """Build variable map for template rendering with version-aware preference."""
    with get_conn() as con:
        df = con.execute(
            """
            SELECT
                rtv.variable_key,
                COALESCE(
                  (
                    SELECT jrv.variable_value
                    FROM job_report_variable_values jrv
                    WHERE jrv.job_id = %s
                      AND jrv.template_id = rtv.template_id
                      AND jrv.variable_key = rtv.variable_key
                      AND (%s IS NULL OR jrv.version_id = %s)
                    ORDER BY
                      CASE
                        WHEN %s IS NOT NULL AND jrv.version_id = %s THEN 0
                        ELSE 1
                      END,
                      jrv.updated_at DESC
                    LIMIT 1
                  ),
                  rtv.default_value
                ) AS value
            FROM report_template_variables rtv
            WHERE rtv.template_id = %s
            ORDER BY rtv.display_order, rtv.variable_key
            """,
            [
                int(job_id),
                version_id,
                version_id,
                version_id,
                version_id,
                int(template_id),
            ],
        ).df()

    values = {}
    if df is None or df.empty:
        return values

    for _, row in df.iterrows():
        key = str(row["variable_key"])
        values[key] = _normalize_render_value(row["value"])
    return values


def _get_template_variable_metadata(template_id: int) -> dict[str, dict[str, object]]:
    """Return metadata map by variable_key for display-friendly DOCX rendering."""
    with get_conn() as con:
        df = con.execute(
            """
            SELECT variable_key, variable_label, section, variable_type, display_order
            FROM report_template_variables
            WHERE template_id = %s
            ORDER BY display_order, variable_key
            """,
            [int(template_id)],
        ).df()

    out: dict[str, dict[str, object]] = {}
    if df is None or df.empty:
        return out

    for _, row in df.iterrows():
        key = str(row.get("variable_key") or "").strip()
        if not key:
            continue
        out[key] = {
            "label": str(row.get("variable_label") or key),
            "section": str(row.get("section") or "General"),
            "type": str(row.get("variable_type") or "text"),
            "order": int(row.get("display_order") or 0),
        }
    return out


def _norm_text(value: object) -> str:
    if value is None:
        return ""
    return str(value).strip().lower()


def _classify_activity_group(scope: str, category: str, report_label: str) -> str:
    """Classify a category into report activity group buckets."""
    scope_norm = _norm_text(scope)
    search_text = f"{_norm_text(category)} {_norm_text(report_label)}"

    commuting_keywords = [
        'commut',
        'employee travel to work',
        'home working',
        'working from home',
        'wfh',
    ]
    business_travel_keywords = [
        'business travel',
        'air travel',
        'rail travel',
        'hotel',
        'taxi',
        'car hire',
        'rental',
        'flight',
    ]
    pgs_keywords = [
        'purchased goods',
        'purchased services',
        'purchased good',
        'capital goods',
        'procurement',
        'supplier',
        'supply chain',
        'spend',
    ]
    energy_keywords = [
        'electricity',
        'gas',
        'fuel',
        'diesel',
        'petrol',
        'heating',
        'cooling',
        'energy',
        'steam',
        'refrigerant',
    ]

    if any(k in search_text for k in commuting_keywords):
        return 'Employee Commuting'

    if any(k in search_text for k in business_travel_keywords):
        return 'Business Travel'

    if any(k in search_text for k in pgs_keywords):
        return 'Purchased Goods & Services (PG&S)'

    if scope_norm in {'scope 1', 'scope 2'} or any(k in search_text for k in energy_keywords):
        return 'Energy'

    return 'Other Emissions'


def _source_family_label(row: dict[str, Any]) -> str:
    source_type = str(row.get("source_type") or "").strip().lower()
    record_type = str(row.get("record_type") or "legacy").strip().lower()
    if source_type == "business_travel":
        return "Business Travel Register"
    if record_type == "source_register":
        return "Asset Register"
    return "Legacy Data Entry"


def _clean_label(value: Any, fallback: str = "Uncategorized") -> str:
    txt = str(value or "").strip()
    if not txt or txt.lower() in {"nan", "none", "null"}:
        return fallback
    return txt


def _is_placeholder_category(value: Any) -> bool:
    txt = str(value or "").strip().lower()
    return not txt or txt in {"nan", "none", "null", "uncategorized", "uncategorised"}


def _dataset_category_label(row: dict[str, Any]) -> str:
    for key in ("dataset_category", "lookup_category", "category", "lookup_level_1", "lookup_level_2", "level_1", "level_2"):
        value = row.get(key)
        if _is_placeholder_category(value):
            continue
        label = _clean_label(value)
        if label:
            return label
    return "Uncategorized"


def _source_factor_id(row: dict[str, Any]) -> str:
    for key in ("original_id", "group_original_id", "factor_db_id", "group_factor_db_id", "dataset_id", "group_dataset_id"):
        value = row.get(key)
        if value not in (None, ""):
            return str(value)
    return ""


def _source_reference_label(row: dict[str, Any]) -> str:
    for key in ("group_name", "asset_identifier", "source_name", "report_label", "category"):
        value = row.get(key)
        if value not in (None, ""):
            return str(value)
    return ""


def _build_activity_grouping(categories: list[dict]) -> tuple[dict[str, list[dict]], dict[str, float], list[dict]]:
    """Return grouped activities, totals by group, and flattened detailed rows."""
    groups: dict[str, list[dict]] = {k: [] for k in ACTIVITY_GROUP_ORDER}
    totals: dict[str, float] = {k: 0.0 for k in ACTIVITY_GROUP_ORDER}

    for cat in categories:
        group = _classify_activity_group(
            str(cat.get('scope', '') or ''),
            str(cat.get('category', '') or ''),
            str(cat.get('report_label', '') or ''),
        )
        emissions = float(cat.get('emissions', 0) or 0)

        enriched = dict(cat)
        enriched['activity_group'] = group
        groups[group].append(enriched)
        totals[group] += emissions

    details: list[dict] = []
    for group in ACTIVITY_GROUP_ORDER:
        sorted_rows = sorted(
            groups[group],
            key=lambda row: float(row.get('emissions', 0) or 0),
            reverse=True,
        )
        for row in sorted_rows:
            qty = float(row.get('qty', 0) or 0)
            uom = row.get('uom') or ''
            source_family = row.get('source_family') or (
                _source_family_label(row)
            )
            factor_id = _source_factor_id(row)
            reference_label = _source_reference_label(row)
            details.append(
                {
                    'activity_group': group,
                    'scope': row.get('scope', ''),
                    'emission_type': row.get('report_label') or row.get('category') or 'Uncategorised',
                    'category': row.get('category', ''),
                    'report_label': row.get('report_label', ''),
                    'emissions': float(row.get('emissions', 0) or 0),
                    'qty': qty,
                    'uom': uom,
                    'source_family': source_family,
                    'record_type': row.get('record_type') or 'legacy',
                    'source_type': row.get('source_type'),
                    'group_name': row.get('group_name'),
                    'source_name': row.get('source_name'),
                    'asset_identifier': row.get('asset_identifier'),
                    'employee_name': row.get('employee_name'),
                    'reference_label': reference_label,
                    'factor_id': factor_id,
                    'factor_db_id': row.get('factor_db_id'),
                    'original_id': row.get('original_id'),
                    'data_source': row.get('data_source') or (f"{uom} Data" if uom else 'Calculated'),
                    'data_confidence': row.get('data_confidence') or ('High' if qty > 0 else 'Medium'),
                }
            )

    return groups, totals, details


def _records_from_df(df):
    if df is None or getattr(df, "empty", True):
        return []
    try:
        return df.where(df.notna(), None).to_dict("records")
    except Exception:
        logger.debug("Failed to load job sites; returning empty list", exc_info=True)
        return []


def _load_legacy_reporting_rows(con, job_id: int) -> list[dict[str, Any]]:
    factor_category_expr = _factor_category_expr(con)
    df = con.execute(
        f"""
        SELECT
            'legacy' AS record_type,
            COALESCE(s.site_name, 'Unassigned') AS site_name,
            jsr.scope,
            CASE
                WHEN COALESCE(TRIM(CAST({factor_category_expr} AS VARCHAR)), '') = ''
                    OR LOWER(TRIM(CAST({factor_category_expr} AS VARCHAR))) IN ('nan', 'none', 'null')
                THEN COALESCE(NULLIF(TRIM(CAST(jsr.category AS VARCHAR)), ''), COALESCE(NULLIF(TRIM(CAST(jsr.level_2 AS VARCHAR)), ''), 'Uncategorized'))
                ELSE TRIM(CAST({factor_category_expr} AS VARCHAR))
            END AS category,
            CASE
                WHEN COALESCE(NULLIF(TRIM(CAST(jsr.report_label AS VARCHAR)), ''), '') <> ''
                THEN TRIM(CAST(jsr.report_label AS VARCHAR))
                ELSE CASE
                    WHEN COALESCE(TRIM(CAST({factor_category_expr} AS VARCHAR)), '') = ''
                        OR LOWER(TRIM(CAST({factor_category_expr} AS VARCHAR))) IN ('nan', 'none', 'null')
                    THEN COALESCE(NULLIF(TRIM(CAST(jsr.category AS VARCHAR)), ''), COALESCE(NULLIF(TRIM(CAST(jsr.level_2 AS VARCHAR)), ''), 'Uncategorized'))
                    ELSE TRIM(CAST({factor_category_expr} AS VARCHAR))
                END
            END AS report_label,
            jsr.dataset_id,
            jsr.factor_db_id,
            jsr.original_id,
            jsr.qty,
            jsr.uom,
            jsr.factor,
            jsr.ghg_unit,
            jsr.apply_pct,
            COALESCE(jsr.data_source, 'Company Data') AS data_source,
            COALESCE(jsr.data_confidence, 'M') AS data_confidence,
            jsr.notes,
            jsr.source_qty,
            jsr.source_uom,
            jsr.month_1, jsr.month_2, jsr.month_3, jsr.month_4, jsr.month_5, jsr.month_6,
            jsr.month_7, jsr.month_8, jsr.month_9, jsr.month_10, jsr.month_11, jsr.month_12,
            jsr.enabled,
            jsr.created_at,
            jsr.updated_at,
            NULL::INTEGER AS source_id,
            NULL::INTEGER AS group_id,
            NULL::VARCHAR AS source_type,
            NULL::VARCHAR AS source_subtype,
            NULL::VARCHAR AS asset_identifier,
            NULL::VARCHAR AS employee_name,
            NULL::VARCHAR AS group_name,
            COALESCE(
                NULLIF(jsr.report_label, ''),
                CASE
                    WHEN COALESCE(TRIM(CAST({factor_category_expr} AS VARCHAR)), '') = ''
                        OR LOWER(TRIM(CAST({factor_category_expr} AS VARCHAR))) IN ('nan', 'none', 'null')
                    THEN COALESCE(NULLIF(TRIM(CAST(jsr.category AS VARCHAR)), ''), COALESCE(NULLIF(TRIM(CAST(jsr.level_2 AS VARCHAR)), ''), 'Uncategorized'))
                    ELSE TRIM(CAST({factor_category_expr} AS VARCHAR))
                END
            ) AS source_name
        FROM job_scope_rows jsr
        LEFT JOIN client_sites s ON s.site_id = jsr.site_id
        LEFT JOIN factor_lookup fl ON fl.db_id = jsr.factor_db_id
        WHERE jsr.job_id = %s AND COALESCE(jsr.enabled, TRUE) = TRUE
        ORDER BY COALESCE(s.site_name, 'Unassigned'), jsr.scope, category, report_label
        """,
        [int(job_id)],
    ).df()
    return _records_from_df(df)


def _load_source_register_rows(con, job_id: int) -> list[dict[str, Any]]:
    factor_category_expr = _factor_category_expr(con)
    df = con.execute(
        f"""
        SELECT
            'source_register' AS record_type,
            COALESCE(cs.site_name, 'Unassigned') AS site_name,
            js.scope,
            CASE
                WHEN COALESCE(TRIM(CAST({factor_category_expr} AS VARCHAR)), '') = ''
                    OR LOWER(TRIM(CAST({factor_category_expr} AS VARCHAR))) IN ('nan', 'none', 'null')
                THEN COALESCE(NULLIF(TRIM(CAST(js.category AS VARCHAR)), ''), 'Uncategorized')
                ELSE TRIM(CAST({factor_category_expr} AS VARCHAR))
            END AS category,
            CASE
                WHEN COALESCE(NULLIF(TRIM(CAST(js.source_name AS VARCHAR)), ''), '') <> ''
                THEN TRIM(CAST(js.source_name AS VARCHAR))
                ELSE COALESCE(
                    NULLIF(TRIM(CAST(g.group_name AS VARCHAR)), ''),
                    CASE
                        WHEN COALESCE(TRIM(CAST({factor_category_expr} AS VARCHAR)), '') = ''
                            OR LOWER(TRIM(CAST({factor_category_expr} AS VARCHAR))) IN ('nan', 'none', 'null')
                        THEN COALESCE(NULLIF(TRIM(CAST(js.category AS VARCHAR)), ''), 'Uncategorized')
                        ELSE TRIM(CAST({factor_category_expr} AS VARCHAR))
                    END
                )
            END AS report_label,
            COALESCE(g.dataset_id, js.dataset_id) AS dataset_id,
            COALESCE(g.factor_db_id, js.factor_db_id) AS factor_db_id,
            COALESCE(g.original_id, js.original_id) AS original_id,
            js.qty,
            COALESCE(g.uom, js.uom) AS uom,
            COALESCE(g.factor, js.factor) AS factor,
            COALESCE(g.ghg_unit, js.ghg_unit) AS ghg_unit,
            js.apply_pct,
            js.calc_tco2e,
            COALESCE(js.data_source, CASE WHEN js.source_type = 'business_travel' THEN 'Business Travel Register' ELSE 'Asset Register' END) AS data_source,
            COALESCE(js.data_confidence, 'M') AS data_confidence,
            js.notes,
            NULL::NUMERIC AS source_qty,
            NULL::VARCHAR AS source_uom,
            NULL::NUMERIC AS month_1,
            NULL::NUMERIC AS month_2,
            NULL::NUMERIC AS month_3,
            NULL::NUMERIC AS month_4,
            NULL::NUMERIC AS month_5,
            NULL::NUMERIC AS month_6,
            NULL::NUMERIC AS month_7,
            NULL::NUMERIC AS month_8,
            NULL::NUMERIC AS month_9,
            NULL::NUMERIC AS month_10,
            NULL::NUMERIC AS month_11,
            NULL::NUMERIC AS month_12,
            js.enabled,
            js.created_at,
            js.updated_at,
            js.source_id,
            js.group_id,
            js.source_type,
            js.source_subtype,
            js.asset_identifier,
            js.employee_name,
            g.group_name,
            g.dataset_id AS group_dataset_id,
            g.factor_db_id AS group_factor_db_id,
            g.original_id AS group_original_id,
            g.factor AS group_factor,
            g.ghg_unit AS group_ghg_unit,
            g.uom AS group_uom,
            js.source_name
        FROM job_emission_sources js
        LEFT JOIN job_emission_groups g ON g.group_id = js.group_id
        LEFT JOIN client_sites cs ON cs.site_id = js.site_id
        LEFT JOIN factor_lookup fl ON fl.db_id = COALESCE(g.factor_db_id, js.factor_db_id)
        WHERE js.job_id = %s AND COALESCE(js.enabled, TRUE) = TRUE
        ORDER BY COALESCE(cs.site_name, 'Unassigned'), js.scope, category, report_label
        """,
        [int(job_id)],
    ).df()
    return _records_from_df(df)


def _load_reporting_rows(con, job_id: int) -> list[dict[str, Any]]:
    return _records_from_df(_load_data_output_rows(con, int(job_id)))


def create_donut_chart(data_dict, colors_dict, title, total_value, center_label="tCO₂e"):
    """
    Create a high-quality donut chart and return it as a base64 encoded PNG.
    Optimized for professional PDF rendering at 25% of page size.
    
    Args:
        data_dict: Dictionary with labels as keys and values as values
        colors_dict: Dictionary with labels as keys and hex colors as values
        title: Chart title
        total_value: Total value to display in center
        center_label: Label to show under the center value
    
    Returns:
        Base64 encoded PNG string
    """
    # Import matplotlib inside function to avoid startup overhead
    import matplotlib
    matplotlib.use('Agg')  # Use non-interactive backend
    import matplotlib.pyplot as plt
    from matplotlib.patches import Circle
    import base64
    
    # High-quality figure settings
    plt.rcParams['figure.dpi'] = 300
    plt.rcParams['savefig.dpi'] = 300
    plt.rcParams['font.family'] = 'sans-serif'
    plt.rcParams['font.sans-serif'] = ['Arial', 'Helvetica', 'DejaVu Sans']
    plt.rcParams['axes.linewidth'] = 0.5

    
    # High-quality figure size (CSS will scale it down)
    fig, ax = plt.subplots(figsize=(4, 4))
    
    # Filter out zero values for the chart
    labels = [k for k, v in data_dict.items() if v > 0]
    values = [v for v in data_dict.values() if v > 0]
    chart_colors = [colors_dict.get(k, '#999999') for k in labels]
    
    if not values:
        # If no data, create a placeholder
        labels = ['No Data']
        values = [1]
        chart_colors = ['#CCCCCC']
    
    # Create labels with scope name, value with thousand separator, and percentage
    labels_with_data = []
    total = sum(values)
    for label, val in zip(labels, values):
        pct = (val / total * 100) if total > 0 else 0
        labels_with_data.append(f'{label}\n{val:,.2f} ({pct:.1f}%)')
    
    # Create donut chart with all info outside (no autopct)
    # pctdistance is not used since we're not using autopct, but we ensure lines are drawn
    wedges, texts = ax.pie(
        values,
        labels=labels_with_data,
        colors=chart_colors,
        startangle=90,
        labeldistance=1.3,
        wedgeprops=dict(width=0.35, edgecolor='white', linewidth=2),
        textprops={'fontsize': 10, 'weight': 'normal', 'ha': 'center'}
    )
    
    # Style the label text (outside with lines)
    for text in texts:
        text.set_fontsize(9)
        text.set_weight('normal')
    
    # Add center circle for donut effect
    centre_circle = Circle((0, 0), 0.60, fc='white', ec='none')
    ax.add_artist(centre_circle)
    
    # Add total in center with thousand separator and doubled font size
    ax.text(0, 0.05, f'{total_value:,.2f}', ha='center', va='center',
            fontsize=32, fontweight='bold', color='#70AD47',
            family='sans-serif')
    ax.text(0, -0.18, center_label, ha='center', va='center',
            fontsize=14, color='#70AD47', family='sans-serif')
    
    # Don't add title here - it will be in HTML
    ax.axis('equal')
    
    # Save to bytes with high quality settings
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight',
                facecolor='white', edgecolor='none',
                pad_inches=0.2, transparent=False)
    plt.close(fig)
    img_buffer.seek(0)
    
    # Encode to base64
    img_base64 = base64.b64encode(img_buffer.read()).decode('utf-8')
    return img_base64


def create_activity_donut(
    activity_totals: dict[str, float],
    colors_dict: dict[str, str],
    title: str = "Emissions by Activity",
    center_label: str = "tCO₂e",
    period_from: object | None = None,
    period_to: object | None = None,
):
    """
    Backward-compatible wrapper for activity donut generation.
    The date inputs are accepted for call-site compatibility but not required.
    """
    chart_data = {
        group: float(activity_totals.get(group, 0.0) or 0.0)
        for group in ACTIVITY_GROUP_ORDER
    }
    total_value = sum(chart_data.values())
    return create_donut_chart(
        chart_data,
        colors_dict,
        title,
        total_value,
        center_label=center_label,
    )


def get_job_data(job_id: int, org_id: str | None = None):
    """Fetch job and client data for the report."""
    org_filter = str(org_id or "").strip() or None
    with get_conn() as con:
        if org_filter:
            job_row = con.execute("""
            SELECT j.job_id, j.client_db_id, j.org_id, j.job_number, j.title, j.reporting_year, j.status,
                   j.reporting_period_start, j.reporting_period_end,
                   c.client_name, c.crm_owner, c.industry, c.logo_url, c.description_long,
                   c.net_zero_year, c.interim_year, c.benchmark_year,
                   c.benchmark_period_start, c.benchmark_period_end,
                   c.interim_s1_pct, c.interim_s2_pct, c.interim_s3_pct,
                   c.target_s1_year, c.target_s2_year, c.target_s3_year,
                   c.target_s1_pct, c.target_s2_pct, c.target_s3_pct,
                   c.addr_city, c.addr_country
            FROM jobs j
            JOIN clients c ON c.db_id = j.client_db_id
            WHERE j.job_id = %s AND c.org_id = %s
        """, [int(job_id), org_filter]).fetchone()
        else:
            job_row = con.execute("""
            SELECT j.job_id, j.client_db_id, j.org_id, j.job_number, j.title, j.reporting_year, j.status,
                   j.reporting_period_start, j.reporting_period_end,
                   c.client_name, c.crm_owner, c.industry, c.logo_url, c.description_long,
                   c.net_zero_year, c.interim_year, c.benchmark_year,
                   c.benchmark_period_start, c.benchmark_period_end,
                   c.interim_s1_pct, c.interim_s2_pct, c.interim_s3_pct,
                   c.target_s1_year, c.target_s2_year, c.target_s3_year,
                   c.target_s1_pct, c.target_s2_pct, c.target_s3_pct,
                   c.addr_city, c.addr_country
            FROM jobs j
            JOIN clients c ON c.db_id = j.client_db_id
            WHERE j.job_id = %s
        """, [int(job_id)]).fetchone()
        
        if not job_row:
            return None
        
        # Fetch milestones
        milestones = con.execute("""
            SELECT data_collection_due, first_draft_due, final_report_due
            FROM job_plan
            WHERE job_id = %s
        """, [int(job_id)]).fetchone()
        
        # Fetch CRP job details for Reporting Elements
        crp_details = con.execute("""
            SELECT num_employees, premises_owned, premises_leased, vehicles_owned, vehicles_leased
            FROM crp_job_details
            WHERE job_id = %s
        """, [int(job_id)]).fetchone()
        
        return {
            'job_id': job_row[0],
            'client_db_id': job_row[1],
            'org_id': job_row[2],
            'job_number': job_row[3],
            'title': job_row[4],
            'reporting_year': job_row[5],
            'status': job_row[6],
            'period_start': job_row[7],
            'period_end': job_row[8],
            'reporting_period_start': job_row[7],
            'reporting_period_end': job_row[8],
            'client_name': job_row[9],
            'crm_owner': job_row[10],
            'industry': job_row[11],
            'logo_url': job_row[12],
            'client_logo_url': job_row[12],  # alias used by interactive_report.html
            'description': job_row[13],
            'net_zero_year': job_row[14],
            'interim_year': job_row[15],
            'benchmark_year': job_row[16],
            'benchmark_period_start': job_row[17],
            'benchmark_period_end': job_row[18],
            'interim_s1_pct': job_row[19],
            'interim_s2_pct': job_row[20],
            'interim_s3_pct': job_row[21],
            'target_s1_year': job_row[22],
            'target_s2_year': job_row[23],
            'target_s3_year': job_row[24],
            'target_s1_pct': job_row[25],
            'target_s2_pct': job_row[26],
            'target_s3_pct': job_row[27],
            'city': job_row[28],
            'country': job_row[29],
            'data_collection_due': milestones[0] if milestones else None,
            'first_draft_due': milestones[1] if milestones else None,
            'final_report_due': milestones[2] if milestones else None,
            # Reporting Elements from crp_job_details
            'no_of_staff': crp_details[0] if crp_details else None,
            'no_premises_owned': crp_details[1] if crp_details else None,
            'no_premises_leased': crp_details[2] if crp_details else None,
            'no_vehicles_owned': crp_details[3] if crp_details else None,
            'no_vehicles_leased': crp_details[4] if crp_details else None,
        }


def _normalize_report_draft_section_key(section_key: str) -> str:
    raw = str(section_key or "").strip().lower()
    normalized = raw.replace(" ", "_").replace("-", "_")
    aliases = {
        "emissions_footprint": "emissions_overview",
        "footprint_summary": "emissions_overview",
        "emissions_overview": "emissions_overview",
        "executive_summary": "executive_summary",
        "actions": "actions",
    }
    return aliases.get(normalized, normalized)


def _coerce_float(value: Any) -> float:
    try:
        if value is None:
            return 0.0
        return float(value)
    except Exception:
        logger.debug("Failed to coerce float in report helper; defaulting to 0.0", exc_info=True)
        return 0.0


def _top_category(categories: list[dict[str, Any]]) -> dict[str, Any] | None:
    if not categories:
        return None
    return max(categories, key=lambda item: _coerce_float(item.get("emissions")))


def _summarize_categories(categories: list[dict[str, Any]], limit: int = 5) -> list[dict[str, Any]]:
    ordered = sorted(categories or [], key=lambda item: _coerce_float(item.get("emissions")), reverse=True)
    summary: list[dict[str, Any]] = []
    for item in ordered[:limit]:
        category = _dataset_category_label(item)
        summary.append(
            {
                "category": category,
                "dataset_category": category,
                "emissions": _coerce_float(item.get("emissions")),
                "scope": str(item.get("scope") or ""),
                "report_label": str(item.get("report_label") or ""),
                "data_source": str(item.get("data_source") or ""),
            }
        )
    return summary


def _build_report_draft_context(job_id: int, template_key: str | None = None) -> dict[str, Any]:
    job_data = get_job_data(job_id)
    if not job_data:
        raise HTTPException(status_code=404, detail="Job not found")

    def _safe_totals() -> dict[str, float]:
        return {"Scope 1": 0.0, "Scope 2": 0.0, "Scope 3": 0.0, "Total": 0.0}

    try:
        scope_totals = get_scope_totals(job_id)
    except Exception:
        logger.warning("Failed to load scope totals for job %s; using zero totals", job_id, exc_info=True)
        scope_totals = _safe_totals()

    try:
        categories = get_emissions_by_category(job_id)
    except Exception:
        logger.warning("Failed to load emissions by category for job %s; using empty list", job_id, exc_info=True)
        categories = []

    try:
        benchmark_totals = get_benchmark_emissions(job_id, job_data.get("benchmark_year"))
    except Exception:
        logger.warning("Failed to load benchmark totals for job %s; using zero totals", job_id, exc_info=True)
        benchmark_totals = _safe_totals()

    try:
        benchmark_job_id = _resolve_benchmark_reference_job(job_id, job_data.get("benchmark_year"))
    except Exception:
        logger.debug("Failed to resolve benchmark reference for job %s; continuing without benchmark", job_id, exc_info=True)
        benchmark_job_id = None

    try:
        previous_job_data = get_job_data(benchmark_job_id) if benchmark_job_id else None
    except Exception:
        logger.debug("Failed to load previous-year job data for job %s; continuing without comparison context", job_id, exc_info=True)
        previous_job_data = None

    try:
        previous_categories = get_emissions_by_category(benchmark_job_id) if benchmark_job_id else []
    except Exception:
        logger.debug("Failed to load previous-year emissions categories for job %s; continuing without comparison data", job_id, exc_info=True)
        previous_categories = []

    try:
        job_actions = get_job_report_actions_payload(job_id)
    except Exception:
        logger.debug("Failed to load report actions payload for job %s; continuing with empty actions", job_id, exc_info=True)
        job_actions = {"items": [], "term_counts": {}}

    try:
        selected_template = _get_job_assigned_template_selection(int(job_id))
    except Exception:
        logger.debug("Failed to load selected report template for job %s; continuing without assignment details", job_id, exc_info=True)
        selected_template = None

    current_total = _coerce_float(scope_totals.get("Total"))
    previous_total = _coerce_float(benchmark_totals.get("Total"))
    change_pct = None
    if previous_total > 0:
        change_pct = ((current_total - previous_total) / previous_total) * 100.0

    top = _top_category(categories)

    context_summary = (
        f"{job_data.get('client_name') or 'Client'} reported {current_total:.1f} tCO₂e in "
        f"{job_data.get('reporting_year') or 'the reporting period'}."
    )
    if change_pct is not None:
        direction = "higher" if change_pct > 0 else "lower" if change_pct < 0 else "flat"
        context_summary += f" Total emissions are {abs(change_pct):.1f}% {direction} than the comparison period."
    if top:
        context_summary += f" Largest category: {top.get('category')} at {_coerce_float(top.get('emissions')):.1f} tCO₂e."

    selected_template_key = selected_template.get("template_key") if selected_template else None

    return {
        "job_id": int(job_id),
        "template_key": template_key or selected_template_key,
        "selected_template": selected_template or {},
        "job_data": job_data,
        "previous_job_data": previous_job_data,
        "scope_totals": scope_totals,
        "benchmark_totals": benchmark_totals,
        "categories": _summarize_categories(categories),
        "previous_categories": _summarize_categories(previous_categories),
        "job_actions": job_actions,
        "context_summary": context_summary,
        "top_category": {
            "category": str(top.get("category") or "") if top else "",
            "emissions": _coerce_float(top.get("emissions")) if top else 0.0,
        }
        if top
        else None,
    }


def _report_draft_template_variable_key(section_key: str) -> str | None:
    normalized = _normalize_report_draft_section_key(section_key)
    mapping = {
        "executive_summary": "executive_summary",
        "emissions_overview": "footprint_summary",
        "actions": "reduction_projects",
    }
    return mapping.get(normalized)


def _serialize_report_draft_row(row: dict[str, Any]) -> dict[str, Any]:
    draft_json = row.get("draft_json")
    parsed_json: dict[str, Any] | None = None
    if isinstance(draft_json, dict):
        parsed_json = draft_json
    elif isinstance(draft_json, str) and draft_json.strip():
        try:
            maybe = json.loads(draft_json)
            if isinstance(maybe, dict):
                parsed_json = maybe
        except Exception:
            logger.debug("Failed to parse draft JSON payload; continuing without structured draft", exc_info=True)
            parsed_json = None

    return {
        "draft_id": int(row.get("draft_id")) if row.get("draft_id") is not None else None,
        "job_id": int(row.get("job_id")) if row.get("job_id") is not None else None,
        "template_id": int(row.get("template_id")) if row.get("template_id") is not None else None,
        "version_id": int(row.get("version_id")) if row.get("version_id") is not None else None,
        "section_key": row.get("section_key"),
        "section_title": row.get("section_title"),
        "draft_text": row.get("draft_text"),
        "draft_json": parsed_json,
        "evidence_hash": row.get("evidence_hash"),
        "provider": row.get("provider"),
        "model": row.get("model"),
        "confidence": row.get("confidence"),
        "status": row.get("status"),
        "revision_number": int(row.get("revision_number")) if row.get("revision_number") is not None else None,
        "created_at": str(row.get("created_at")) if row.get("created_at") else None,
        "updated_at": str(row.get("updated_at")) if row.get("updated_at") else None,
        "created_by": row.get("created_by"),
        "updated_by": row.get("updated_by"),
    }


def _upsert_report_draft(
    *,
    con,
    job_id: int,
    template_id: int,
    version_id: int,
    section_key: str,
    section_title: str | None,
    draft_text: str,
    draft_json: dict[str, Any] | None,
    evidence_hash: str | None,
    provider: str | None,
    model: str | None,
    confidence: str | None,
    status: str | None,
    actor: str,
) -> dict[str, Any]:
    _ensure_report_drafts_schema(con)

    normalized_key = _normalize_report_draft_section_key(section_key)
    draft_text_value = str(draft_text or "").strip()
    draft_json_text = json.dumps(draft_json or {}, ensure_ascii=False, sort_keys=True, default=str)
    normalized_status = str(status or "draft").strip().lower() or "draft"
    job_data = get_job_data(int(job_id))
    org_id = str(job_data.get("org_id") or "").strip() if job_data else ""
    current_row = con.execute(
        """
        SELECT draft_id, revision_number
        FROM job_report_drafts
        WHERE job_id = %s AND template_id = %s AND version_id = %s AND section_key = %s
        """,
        [int(job_id), int(template_id), int(version_id), normalized_key],
    ).fetchone()

    if current_row:
        next_revision = int(current_row[1] or 0) + 1
        con.execute(
            """
            UPDATE job_report_drafts
            SET section_title = %s,
                draft_text = %s,
                draft_json = %s,
                evidence_hash = %s,
                provider = %s,
                model = %s,
                confidence = %s,
                status = %s,
                revision_number = %s,
                updated_at = NOW(),
                updated_by = %s
            WHERE draft_id = %s
            """,
            [
                section_title,
                draft_text_value,
                draft_json_text,
                evidence_hash,
                provider,
                model,
                confidence,
                normalized_status,
                next_revision,
                actor,
                int(current_row[0]),
            ],
        )
    else:
        con.execute(
            """
            INSERT INTO job_report_drafts (
              org_id, job_id, template_id, version_id, section_key, section_title, draft_text, draft_json,
              evidence_hash, provider, model, confidence, status, revision_number, created_by, updated_by
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 1, %s, %s)
            """,
            [
                org_id or None,
                int(job_id),
                int(template_id),
                int(version_id),
                normalized_key,
                section_title,
                draft_text_value,
                draft_json_text,
                evidence_hash,
                provider,
                model,
                confidence,
                normalized_status,
                actor,
                actor,
            ],
        )

    variable_key = _report_draft_template_variable_key(normalized_key)
    if variable_key:
        con.execute(
            """
            INSERT INTO job_report_variable_values (
              org_id, job_id, template_id, version_id, variable_key, variable_value, updated_at, updated_by
            )
            VALUES (%s, %s, %s, %s, %s, %s, NOW(), %s)
            ON CONFLICT (job_id, template_id, version_id, variable_key)
            DO UPDATE SET
              org_id = EXCLUDED.org_id,
              variable_value = EXCLUDED.variable_value,
              updated_at = NOW(),
              updated_by = EXCLUDED.updated_by
            """,
            [
                org_id or None,
                int(job_id),
                int(template_id),
                int(version_id),
                variable_key,
                draft_text_value,
                actor,
            ],
        )

    saved_df = con.execute(
        """
        SELECT *
        FROM job_report_drafts
        WHERE job_id = %s AND template_id = %s AND version_id = %s AND section_key = %s
        """,
        [int(job_id), int(template_id), int(version_id), normalized_key],
    ).df()
    if saved_df is not None and not saved_df.empty:
        return _serialize_report_draft_row(saved_df.iloc[0].to_dict())
    return {
        "job_id": int(job_id),
        "template_id": int(template_id),
        "version_id": int(version_id),
        "section_key": normalized_key,
        "section_title": section_title,
        "draft_text": draft_text_value,
        "draft_json": draft_json or {},
        "evidence_hash": evidence_hash,
        "provider": provider,
        "model": model,
        "confidence": confidence,
        "status": normalized_status,
    }


def _delete_report_drafts_for_missing_sections(
    *,
    con,
    job_id: int,
    template_id: int,
    version_id: int,
    keep_section_keys: set[str],
) -> None:
    keep_section_keys = {_normalize_report_draft_section_key(key) for key in keep_section_keys}
    existing_rows = con.execute(
        """
        SELECT section_key
        FROM job_report_drafts
        WHERE job_id = %s AND template_id = %s AND version_id = %s
        """,
        [int(job_id), int(template_id), int(version_id)],
    ).fetchall()
    existing_keys = {str(row[0]) for row in existing_rows or [] if row and row[0]}
    remove_keys = sorted(existing_keys - keep_section_keys)
    if not remove_keys:
        return

    placeholders = ", ".join(["%s"] * len(remove_keys))
    con.execute(
        f"""
        DELETE FROM job_report_drafts
        WHERE job_id = %s AND template_id = %s AND version_id = %s AND section_key IN ({placeholders})
        """,
        [int(job_id), int(template_id), int(version_id), *remove_keys],
    )

    variable_keys_to_remove = [
        _report_draft_template_variable_key(section_key)
        for section_key in remove_keys
        if _report_draft_template_variable_key(section_key)
    ]
    if variable_keys_to_remove:
        var_placeholders = ", ".join(["%s"] * len(variable_keys_to_remove))
        con.execute(
            f"""
            DELETE FROM job_report_variable_values
            WHERE job_id = %s AND template_id = %s AND version_id = %s AND variable_key IN ({var_placeholders})
            """,
            [int(job_id), int(template_id), int(version_id), *variable_keys_to_remove],
        )


def _load_report_drafts(con, job_id: int, template_id: int, version_id: int) -> list[dict[str, Any]]:
    _ensure_report_drafts_schema(con)
    rows = con.execute(
        """
        SELECT *
        FROM job_report_drafts
        WHERE job_id = %s AND template_id = %s AND version_id = %s
        ORDER BY section_key ASC, updated_at DESC, draft_id DESC
        """,
        [int(job_id), int(template_id), int(version_id)],
    ).df()
    if rows is None or rows.empty:
        return []
    return [_serialize_report_draft_row(row.to_dict()) for _, row in rows.iterrows()]


def _resolve_draft_section_inputs(
    sections: list[ReportDraftSectionPayload],
) -> list[ReportDraftSectionPayload]:
    return [item for item in sections if isinstance(item, ReportDraftSectionPayload)]


def get_scope_totals(job_id: int):
    """Get emissions totals by scope."""
    with get_conn() as con:
        resolver = JobMonthlyEmissionsResolver(con, int(job_id))
        rows = _load_reporting_rows(con, int(job_id))
        if not rows:
            return {"Scope 1": 0.0, "Scope 2": 0.0, "Scope 3": 0.0, "Total": 0.0}
        data_df = _load_data_output_rows(con, int(job_id))
        if data_df is None or data_df.empty:
            return {"Scope 1": 0.0, "Scope 2": 0.0, "Scope 3": 0.0, "Total": 0.0}
        _, totals = _build_scope_summary(data_df, resolver)
        return totals


def get_emissions_by_category(job_id: int):
    """Get emissions breakdown by category."""
    with get_conn() as con:
        try:
            resolver = JobMonthlyEmissionsResolver(con, int(job_id))
            rows = _load_reporting_rows(con, int(job_id))
        except Exception:
            logger.warning("Failed to load reporting rows for job %s; returning empty emissions breakdown", job_id, exc_info=True)
            return []

        if not rows:
            return []

        categories = []
        for row in rows:
            try:
                metrics = combined_row_metrics(row, resolver)
                qty_val = float(metrics.get("display_qty") or 0.0)
                emissions = float(metrics.get("calc_tco2e") or 0.0)
            except Exception:
                logger.debug("Skipping malformed emissions row for job %s", job_id, exc_info=True)
                continue

            scope_val = row.get('scope', '')
            cat_val = _dataset_category_label(row)
            report_label_val = row.get('report_label', '')
            activity_group_val = _classify_activity_group(
                str(scope_val), str(cat_val), str(report_label_val)
            )
            categories.append({
                'scope': scope_val,
                'category': cat_val,
                'dataset_category': cat_val,
                'report_label': report_label_val,
                'activity_group': activity_group_val,
                'qty': qty_val,
                'uom': metrics.get('display_uom') or '',
                'emissions': emissions,
                'data_source': row.get('data_source') or '',
                'data_confidence': row.get('data_confidence') or '',
                'record_type': row.get('record_type') or 'legacy',
                'source_type': row.get('source_type'),
                'source_family': row.get('source_family') or _source_family_label(row),
                'group_name': row.get('group_name'),
                'source_name': row.get('source_name'),
                'asset_identifier': row.get('asset_identifier'),
                'employee_name': row.get('employee_name'),
                'original_id': row.get('original_id'),
                'factor_id': _source_factor_id(row),
                'reference_label': _source_reference_label(row),
                'factor_db_id': row.get('factor_db_id'),
            })

        return categories


def _resolve_benchmark_reference_job(job_id: int, benchmark_year: int | None) -> int | None:
    """
    Resolve benchmark comparison job for a given job.
    Priority:
    1) Explicit benchmark year job (excluding current), prefer is_benchmark then latest period.
    2) Latest prior reporting period for same client.
    3) Prior reporting year fallback.
    """
    with get_conn() as con:
        current = con.execute(
            """
            SELECT job_id, client_db_id, reporting_year, reporting_period_end
            FROM jobs
            WHERE job_id = %s
            """,
            [int(job_id)],
        ).fetchone()
        if not current:
            return None

        cur_job_id = int(current[0])
        client_db_id = int(current[1])
        cur_reporting_year = int(current[2]) if current[2] is not None else None
        cur_period_end = current[3]

        if benchmark_year is not None:
            by_year = con.execute(
                """
                SELECT job_id
                FROM jobs
                WHERE client_db_id = %s
                  AND reporting_year = %s
                  AND job_id <> %s
                ORDER BY COALESCE(is_benchmark, FALSE) DESC,
                         reporting_period_end DESC NULLS LAST,
                         job_id DESC
                LIMIT 1
                """,
                [client_db_id, int(benchmark_year), cur_job_id],
            ).fetchone()
            if by_year and by_year[0] is not None:
                return int(by_year[0])

        if cur_period_end is not None:
            prior_period = con.execute(
                """
                SELECT job_id
                FROM jobs
                WHERE client_db_id = %s
                  AND reporting_period_end IS NOT NULL
                  AND reporting_period_end < %s
                  AND job_id <> %s
                ORDER BY reporting_period_end DESC, job_id DESC
                LIMIT 1
                """,
                [client_db_id, cur_period_end, cur_job_id],
            ).fetchone()
            if prior_period and prior_period[0] is not None:
                return int(prior_period[0])

        if cur_reporting_year is not None:
            prior_year = con.execute(
                """
                SELECT job_id
                FROM jobs
                WHERE client_db_id = %s
                  AND reporting_year IS NOT NULL
                  AND reporting_year < %s
                  AND job_id <> %s
                ORDER BY reporting_year DESC, reporting_period_end DESC NULLS LAST, job_id DESC
                LIMIT 1
                """,
                [client_db_id, cur_reporting_year, cur_job_id],
            ).fetchone()
            if prior_year and prior_year[0] is not None:
                return int(prior_year[0])

        return None


def get_benchmark_emissions(job_id: int, benchmark_year: int | None):
    """Get benchmark emissions for comparison using period-aware benchmark resolution."""
    with get_conn() as con:
        ensure_client_benchmark_columns(con)
        job_row = con.execute(
            "SELECT client_db_id FROM jobs WHERE job_id = %s",
            [int(job_id)],
        ).fetchone()
        if job_row and job_row[0] is not None:
            client_benchmark = get_client_benchmark_metrics(con, int(job_row[0]))
            if client_benchmark:
                return {
                    "Scope 1": float(client_benchmark.get("scope1") or 0),
                    "Scope 2": float(client_benchmark.get("scope2") or 0),
                    "Scope 3": float(client_benchmark.get("scope3") or 0),
                    "Total": float(client_benchmark.get("total") or 0),
                }

    benchmark_job_id = _resolve_benchmark_reference_job(int(job_id), benchmark_year)
    if benchmark_job_id is None:
        return {'Scope 1': 0, 'Scope 2': 0, 'Scope 3': 0, 'Total': 0}
    return get_scope_totals(int(benchmark_job_id))


def get_intensity_metrics(job_id: int):
    """Get intensity metrics for the job.
    
    Reads from jobs.intensity_metrics JSONB column (where frontend saves data).
    Falls back to job_report_metadata or crp_job_details for employee number.
    """
    import logging
    logger = logging.getLogger(__name__)
    
    try:
        with get_conn() as con:
            # Get intensity metrics from jobs.intensity_metrics JSONB column
            job_row = con.execute("""
                SELECT intensity_metrics 
                FROM jobs 
                WHERE job_id = %s
            """, [int(job_id)]).fetchone()
            
            result = []
            
            # Get total emissions for intensity calculation
            scope_totals = get_scope_totals(int(job_id))
            total_emissions = scope_totals.get('Total', 0)
            
            logger.warning(f"[INTENSITY DEBUG] Job {job_id} intensity_metrics raw: {job_row[0] if job_row else 'None'}")
            
            # Parse the intensity_metrics JSON from the jobs table
            if job_row and job_row[0]:
                metrics_dict = job_row[0]
                logger.warning(f"[INTENSITY DEBUG] metrics_dict type: {type(metrics_dict)}, content: {metrics_dict}")
                
                if isinstance(metrics_dict, dict):
                    for key, metric in metrics_dict.items():
                        logger.warning(f"[INTENSITY DEBUG] Processing metric key={key}, metric={metric}")
                        
                        if isinstance(metric, dict):
                            metric_value = metric.get('value', 0) or 0
                            metric_label = metric.get('label', key)
                            divider = metric.get('divider', 1) or 1
                            
                            logger.warning(f"[INTENSITY DEBUG] Adding metric: name={metric_label}, value={metric_value}, divider={divider}")
                            
                            # Calculate emissions proportional to this metric's share
                            # For now, use total emissions as the emissions for each metric
                            # (a more accurate approach would calculate per-scope emissions)
                            
                            result.append({
                                'name': metric_label,
                                'value': metric_value,
                                'divider': divider,
                                'unit': key,
                                'emissions': total_emissions
                            })
            
            logger.warning(f"[INTENSITY DEBUG] Final result: {result}")
            
            # If no metrics found in jobs table, try legacy sources for employee number
            if not result:
                employee_number = None
                
                # Try job_report_metadata first
                try:
                    row = con.execute("""
                        SELECT employee_number 
                        FROM job_report_metadata 
                        WHERE job_id = %s
                    """, [int(job_id)]).fetchone()
                    if row and row[0] is not None:
                        employee_number = float(row[0])
                except Exception:
                    logger.debug("Legacy employee_number lookup from job_report_metadata failed; trying fallback source", exc_info=True)
                
                # Fall back to crp_job_details
                if employee_number is None:
                    try:
                        row = con.execute("""
                            SELECT num_employees 
                            FROM crp_job_details 
                            WHERE job_id = %s
                        """, [int(job_id)]).fetchone()
                        if row and row[0] is not None:
                            employee_number = float(row[0])
                    except Exception:
                        logger.debug("Legacy employee_number lookup from crp_job_details failed; continuing without FTE metric", exc_info=True)
                
                # If we have an employee number, add it as FTE metric
                if employee_number is not None and employee_number > 0:
                    result.append({
                        'name': 'Employee',
                        'value': employee_number,
                        'unit': 'employees',
                        'emissions': total_emissions
                    })
            
            return result
    except Exception:
        logger.debug("Failed to build intensity metrics; returning empty list", exc_info=True)
        return []


def get_job_sites(job_id: int):
    """Get sites for a job's client for the Reporting Sites table."""
    try:
        with get_conn() as con:
            # Get client_db_id from job
            job_row = con.execute(
                "SELECT client_db_id FROM jobs WHERE job_id = %s",
                [int(job_id)]
            ).fetchone()
            
            if not job_row or not job_row[0]:
                return []
            
            client_db_id = job_row[0]
            
            # Get sites for this client (active and recently vacated)
            sites = con.execute("""
                SELECT site_name, location, is_registered_office, vacated_date
                FROM client_sites
                WHERE client_db_id = %s 
                  AND (archived = FALSE OR archived IS NULL)
                  AND (vacated_date IS NULL OR vacated_date >= CURRENT_DATE - INTERVAL '2 years')
                ORDER BY is_registered_office DESC, site_name
            """, [int(client_db_id)]).fetchall()
            
            result = []
            for site in sites:
                result.append({
                    'site_name': site[0],
                    'location': site[1],
                    'is_registered_office': site[2],
                    'vacated_date': site[3],
                })
            return result
    except Exception:
        logger.debug("Failed to load job sites for report; returning empty list", exc_info=True)
        return []


def get_emissions_by_site(job_id: int):
    """Get emissions breakdown by site."""
    try:
        with get_conn() as con:
            resolver = JobMonthlyEmissionsResolver(con, int(job_id))
            rows = _load_reporting_rows(con, int(job_id))

            if not rows:
                return {}
            
            sites = {}
            for row in rows:
                site_name = row.get('site_name', 'Unassigned')
                scope = row.get('scope', '')
                emissions = float(combined_row_metrics(row, resolver).get("calc_tco2e") or 0.0)
                
                if site_name not in sites:
                    sites[site_name] = {'Scope 1': 0, 'Scope 2': 0, 'Scope 3': 0, 'Total': 0}
                
                if scope in sites[site_name]:
                    sites[site_name][scope] += emissions
                    sites[site_name]['Total'] += emissions
            
            return sites
    except Exception:
        logger.debug("Failed to load emissions by site; returning empty dict", exc_info=True)
        return {}


def get_site_emissions_breakdowns(job_id: int) -> dict[str, Any]:
    """Build site-level emissions views for report sections and appendix.

    Returns empty/hidden structures when there is 0-1 site in the reporting data.
    """
    with get_conn() as con:
        job_row = con.execute(
            "SELECT client_db_id FROM jobs WHERE job_id = %s",
            [int(job_id)],
        ).fetchone()
        client_db_id = int(job_row[0]) if job_row and job_row[0] is not None else None

        configured_sites_df = None
        if client_db_id is not None:
            configured_sites_df = con.execute(
                """
                SELECT site_name
                FROM client_sites
                WHERE client_db_id = %s
                  AND (archived = FALSE OR archived IS NULL)
                  AND (vacated_date IS NULL OR vacated_date >= CURRENT_DATE - INTERVAL '2 years')
                ORDER BY site_name
                """,
                [int(client_db_id)],
            ).df()

        resolver = JobMonthlyEmissionsResolver(con, int(job_id))
        rows = _load_reporting_rows(con, int(job_id))

        empty_result = {
            "show_site_tables": False,
            "show_appendix": False,
            "site_count": 0,
            "overall": [],
            "scope": [],
            "activity": [],
            "appendix_rows": [],
        }
        configured_site_names: list[str] = []
        if configured_sites_df is not None and not configured_sites_df.empty:
            configured_site_names = [
                str(r.get("site_name") or "").strip()
                for _, r in configured_sites_df.iterrows()
                if str(r.get("site_name") or "").strip()
            ]

        if not rows:
            # No emissions rows yet, but if client has multiple sites configured
            # still show empty site tables to support the reporting structure.
            if len(configured_site_names) > 1:
                zero_scope_rows = [
                    {
                        "site_name": s,
                        "scope_1": 0.0,
                        "scope_2": 0.0,
                        "scope_3": 0.0,
                        "total": 0.0,
                    }
                    for s in configured_site_names
                ]
                return {
                    "show_site_tables": True,
                    "show_appendix": False,
                    "site_count": len(configured_site_names),
                    "overall": [{"site_name": s, "total": 0.0, "pct_total": 0.0} for s in configured_site_names],
                    "scope": zero_scope_rows,
                    "activity": [
                        {
                            "site_name": s,
                            "energy": 0.0,
                            "business_travel": 0.0,
                            "employee_commuting": 0.0,
                            "pgs": 0.0,
                            "other": 0.0,
                            "total": 0.0,
                        }
                        for s in configured_site_names
                    ],
                    "appendix_rows": [],
                }
            return empty_result

        site_scope_totals: dict[str, dict[str, float]] = {}
        site_activity_totals: dict[str, dict[str, float]] = {}
        appendix_rows: list[dict[str, Any]] = []

        # Seed with configured sites so zero-emission sites are still shown.
        for s in configured_site_names:
            if s not in site_scope_totals:
                site_scope_totals[s] = {"Scope 1": 0.0, "Scope 2": 0.0, "Scope 3": 0.0, "Total": 0.0}
            if s not in site_activity_totals:
                site_activity_totals[s] = {k: 0.0 for k in ACTIVITY_GROUP_ORDER}
                site_activity_totals[s]["Total"] = 0.0

        for row in rows:
            site_name = str(row.get("site_name") or "Unassigned")
            scope = str(row.get("scope") or "")
            category = _dataset_category_label(row)
            report_label = str(row.get("report_label") or category)

            metrics = combined_row_metrics(row, resolver)
            qty_val = float(metrics.get("display_qty") or 0.0)
            uom = str(metrics.get("display_uom") or "")
            emissions = float(metrics.get("calc_tco2e") or 0.0)

            if site_name not in site_scope_totals:
                site_scope_totals[site_name] = {"Scope 1": 0.0, "Scope 2": 0.0, "Scope 3": 0.0, "Total": 0.0}
            if scope in ("Scope 1", "Scope 2", "Scope 3"):
                site_scope_totals[site_name][scope] += emissions
            site_scope_totals[site_name]["Total"] += emissions

            group = _classify_activity_group(scope, category, report_label)
            if site_name not in site_activity_totals:
                site_activity_totals[site_name] = {k: 0.0 for k in ACTIVITY_GROUP_ORDER}
                site_activity_totals[site_name]["Total"] = 0.0
            site_activity_totals[site_name][group] += emissions
            site_activity_totals[site_name]["Total"] += emissions

            appendix_rows.append(
                {
                    "site_name": site_name,
                    "scope": scope,
                    "activity_group": group,
                    "emission_type": report_label,
                    "category": category,
                    "dataset_category": category,
                    "record_type": row.get("record_type") or "legacy",
                    "data_source": row.get("data_source") or "",
                    "data_confidence": row.get("data_confidence") or "",
                    "source_family": row.get("source_family") or (
                        _source_family_label(row)
                    ),
                    "source_type": row.get("source_type"),
                    "source_name": row.get("source_name") or report_label,
                    "group_name": row.get("group_name"),
                    "asset_identifier": row.get("asset_identifier"),
                    "employee_name": row.get("employee_name"),
                    "original_id": row.get("original_id"),
                    "factor_id": _source_factor_id(row),
                    "reference_label": _source_reference_label(row),
                    "factor_db_id": row.get("factor_db_id"),
                    "qty": qty_val,
                    "uom": uom,
                    "emissions": emissions,
                }
            )

        # Keep only non-empty sites and deterministic ordering.
        scope_rows = [
            {
                "site_name": site,
                "scope_1": round(vals.get("Scope 1", 0.0), 2),
                "scope_2": round(vals.get("Scope 2", 0.0), 2),
                "scope_3": round(vals.get("Scope 3", 0.0), 2),
                "total": round(vals.get("Total", 0.0), 2),
            }
            for site, vals in site_scope_totals.items()
        ]
        scope_rows.sort(key=lambda r: (-float(r["total"]), str(r["site_name"]).lower()))

        grand_total = sum(float(r["total"]) for r in scope_rows)
        overall_rows = [
            {
                "site_name": r["site_name"],
                "total": r["total"],
                "pct_total": (round((float(r["total"]) / grand_total) * 100.0, 1) if grand_total > 0 else 0.0),
            }
            for r in scope_rows
        ]

        activity_rows = []
        for site_row in scope_rows:
            site = str(site_row["site_name"])
            vals = site_activity_totals.get(site, {})
            activity_rows.append(
                {
                    "site_name": site,
                    "energy": round(float(vals.get("Energy", 0.0) or 0.0), 2),
                    "business_travel": round(float(vals.get("Business Travel", 0.0) or 0.0), 2),
                    "employee_commuting": round(float(vals.get("Employee Commuting", 0.0) or 0.0), 2),
                    "pgs": round(float(vals.get("Purchased Goods & Services (PG&S)", 0.0) or 0.0), 2),
                    "other": round(float(vals.get("Other Emissions", 0.0) or 0.0), 2),
                    "total": round(float(vals.get("Total", 0.0) or 0.0), 2),
                }
            )

        appendix_rows = sorted(
            appendix_rows,
            key=lambda r: (
                str(r.get("site_name") or "").lower(),
                str(r.get("scope") or "").lower(),
                str(r.get("activity_group") or "").lower(),
                str(r.get("emission_type") or "").lower(),
            ),
        )

        site_count = len(scope_rows)
        if site_count <= 1:
            return {
                "show_site_tables": False,
                "show_appendix": bool(appendix_rows),
                "site_count": site_count,
                "overall": overall_rows,
                "scope": scope_rows,
                "activity": activity_rows,
                "appendix_rows": appendix_rows,
            }

        return {
            "show_site_tables": True,
            "show_appendix": bool(appendix_rows),
            "site_count": site_count,
            "overall": overall_rows,
            "scope": scope_rows,
            "activity": activity_rows,
            "appendix_rows": appendix_rows,
        }


def _get_benchmark_job_id(job_id: int, benchmark_year: int | None) -> int | None:
    return _resolve_benchmark_reference_job(int(job_id), benchmark_year)


def _safe_pct_change(current: float, benchmark: float) -> float | None:
    if benchmark == 0:
        return None
    return round(((current - benchmark) / benchmark) * 100.0, 1)


def _build_scope_comparison(scope_totals: dict[str, float], benchmark_totals: dict[str, float]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for label in ("Scope 1", "Scope 2", "Scope 3", "Total"):
        cur = float(scope_totals.get(label, 0) or 0)
        ben = float(benchmark_totals.get(label, 0) or 0)
        rows.append(
            {
                "label": label,
                "current": round(cur, 2),
                "benchmark": round(ben, 2),
                "change_pct": _safe_pct_change(cur, ben),
            }
        )
    return rows


def _build_activity_comparison(activity_totals: dict[str, float], benchmark_activity_totals: dict[str, float]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for group in ACTIVITY_GROUP_ORDER:
        cur = float(activity_totals.get(group, 0) or 0)
        ben = float(benchmark_activity_totals.get(group, 0) or 0)
        rows.append(
            {
                "group": group,
                "current": round(cur, 2),
                "benchmark": round(ben, 2),
                "change_pct": _safe_pct_change(cur, ben),
            }
        )
    total_cur = sum(float(activity_totals.get(g, 0) or 0) for g in ACTIVITY_GROUP_ORDER)
    total_ben = sum(float(benchmark_activity_totals.get(g, 0) or 0) for g in ACTIVITY_GROUP_ORDER)
    rows.append(
        {
            "group": "Total",
            "current": round(total_cur, 2),
            "benchmark": round(total_ben, 2),
            "change_pct": _safe_pct_change(total_cur, total_ben),
        }
    )
    return rows


def _build_site_overall_comparison(current_breakdowns: dict[str, Any], benchmark_breakdowns: dict[str, Any]) -> list[dict[str, Any]]:
    current_rows = current_breakdowns.get("overall", []) if current_breakdowns else []
    benchmark_rows = benchmark_breakdowns.get("overall", []) if benchmark_breakdowns else []
    current_map = {str(r.get("site_name") or ""): float(r.get("total") or 0) for r in current_rows}
    benchmark_map = {str(r.get("site_name") or ""): float(r.get("total") or 0) for r in benchmark_rows}
    site_names = sorted(set(current_map.keys()) | set(benchmark_map.keys()))
    rows: list[dict[str, Any]] = []
    for site in site_names:
        cur = current_map.get(site, 0.0)
        ben = benchmark_map.get(site, 0.0)
        rows.append(
            {
                "site_name": site,
                "current": round(cur, 2),
                "benchmark": round(ben, 2),
                "change_pct": _safe_pct_change(cur, ben),
            }
        )
    rows.sort(key=lambda r: (-float(r["current"]), str(r["site_name"]).lower()))
    return rows


def get_job_target_data(job_id: int) -> dict:
    """
    Get target data for a job including baseline year, net zero target, and interim targets.
    Falls back to client-level data if job-specific data is not set.
    """
    with get_conn() as con:
        # First try to get job-specific data
        job_row = con.execute("""
            SELECT j.baseline_year, j.net_zero_target_year, j.glossary_terms, j.interim_target_year,
                   c.net_zero_year, c.interim_year,
                   c.interim_s1_pct, c.interim_s2_pct, c.interim_s3_pct,
                   c.target_s1_year, c.target_s2_year, c.target_s3_year,
                   c.target_s1_pct, c.target_s2_pct, c.target_s3_pct,
                   c.net_zero_target_reduction_pct
            FROM jobs j
            LEFT JOIN clients c ON c.db_id = j.client_db_id
            WHERE j.job_id = %s
        """, [int(job_id)]).fetchone()

        if not job_row:
            return {}

        # Job-specific values (may be None)
        baseline_year = job_row[0]
        net_zero_target_year = job_row[1]
        glossary_terms = job_row[2]
        interim_target_year = job_row[3]

        # Client-level fallback values
        client_net_zero_year = job_row[4] or 2050
        client_interim_year = job_row[5] or 2035

        # Use job-specific if set, otherwise fall back to client values
        if net_zero_target_year is None:
            net_zero_target_year = client_net_zero_year
        if interim_target_year is None:
            interim_target_year = client_interim_year
        if baseline_year is None:
            baseline_year = None

        # Per-scope interim percentages
        interim_s1_pct = job_row[6] or 50
        interim_s2_pct = job_row[7] or 50
        interim_s3_pct = job_row[8] or 50

        # Net Zero target reduction % from client (e.g. 90 means reduce by 90%)
        net_zero_target_reduction_pct = job_row[15] or 90

        return {
            'baseline_year': baseline_year,
            'net_zero_target_year': net_zero_target_year,
            'glossary_terms': glossary_terms,
            'interim_target_year': interim_target_year,
            'interim_year': interim_target_year,
            'interim_pct': interim_s1_pct,
            'interim_s1_pct': interim_s1_pct,
            'interim_s2_pct': interim_s2_pct,
            'interim_s3_pct': interim_s3_pct,
            'target_pct': net_zero_target_reduction_pct,
            'net_zero_target_reduction_pct': net_zero_target_reduction_pct,
        }


def _parse_legacy_glossary_terms(raw: Any) -> list[dict[str, str]]:
    if raw is None:
        return []
    data = raw
    if isinstance(raw, str):
        txt = raw.strip()
        if not txt:
            return []
        try:
            data = json.loads(txt)
        except Exception:
            logger.debug("Failed to parse legacy glossary payload; returning empty list", exc_info=True)
            return []
    if not isinstance(data, list):
        return []

    out: list[dict[str, str]] = []
    for item in data:
        if isinstance(item, dict):
            term = str(item.get("term") or item.get("title") or "").strip()
            definition = str(item.get("definition") or item.get("content") or "").strip()
            if term and definition:
                out.append({"term": term, "definition": definition})
    return out


def _ensure_glossary_cards_and_fetch(job_id: int, job_data: dict[str, Any], legacy_glossary_terms: Any) -> list[dict[str, str]]:
    """Ensure glossary cards exist in Data Bank, then return active glossary terms for report rendering."""
    global _glossary_schema_seeded
    with get_conn() as con:
        # DDL + seed inserts run only once per process — avoids table-lock overhead on every page load.
        if not _glossary_schema_seeded:
            con.execute(
                """
                CREATE TABLE IF NOT EXISTS databank_subjects_lookup (
                  subject_id INTEGER GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                  name VARCHAR NOT NULL UNIQUE,
                  is_active BOOLEAN NOT NULL DEFAULT TRUE,
                  created_at TIMESTAMP NOT NULL DEFAULT NOW()
                )
                """
            )
            con.execute(
                """
                CREATE TABLE IF NOT EXISTS databank_cards (
                  card_id INTEGER GENERATED BY DEFAULT AS IDENTITY PRIMARY KEY,
                  subject_id INTEGER REFERENCES databank_subjects_lookup(subject_id),
                  category VARCHAR NOT NULL,
                  country VARCHAR,
                  reporting_year INTEGER,
                  title TEXT NOT NULL,
                  content TEXT NOT NULL,
                  source_type VARCHAR,
                  source_url TEXT,
                  reference_text TEXT,
                  tags TEXT,
                  created_by VARCHAR NOT NULL,
                  created_at TIMESTAMP NOT NULL DEFAULT NOW(),
                  updated_at TIMESTAMP,
                  is_active BOOLEAN NOT NULL DEFAULT TRUE
                )
                """
            )
            con.execute("ALTER TABLE databank_cards ADD COLUMN IF NOT EXISTS tags TEXT")
            con.execute("ALTER TABLE databank_cards ADD COLUMN IF NOT EXISTS source_type VARCHAR")
            con.execute("ALTER TABLE databank_cards ADD COLUMN IF NOT EXISTS source_url TEXT")
            con.execute("ALTER TABLE databank_cards ADD COLUMN IF NOT EXISTS reference_text TEXT")
            con.execute("ALTER TABLE databank_cards ADD COLUMN IF NOT EXISTS updated_at TIMESTAMP")
            con.execute("ALTER TABLE databank_cards ADD COLUMN IF NOT EXISTS is_active BOOLEAN NOT NULL DEFAULT TRUE")

            con.execute(
                """
                INSERT INTO databank_subjects_lookup (name, is_active)
                SELECT 'Glossary', TRUE
                WHERE NOT EXISTS (
                  SELECT 1 FROM databank_subjects_lookup WHERE lower(name) = 'glossary'
                )
                """
            )

            subject_row_seed = con.execute(
                "SELECT subject_id FROM databank_subjects_lookup WHERE lower(name) = 'glossary' LIMIT 1"
            ).fetchone()
            if subject_row_seed:
                seed_subject_id = int(subject_row_seed[0])
                seed_entries = list(DEFAULT_GLOSSARY_ENTRIES)
                seed_entries.extend(_parse_legacy_glossary_terms(legacy_glossary_terms))
                seen: set[str] = set()
                for entry in seed_entries:
                    term = str(entry.get("term") or "").strip()
                    definition = str(entry.get("definition") or "").strip()
                    key = term.lower()
                    if not term or not definition or key in seen:
                        continue
                    seen.add(key)
                    con.execute(
                        """
                        INSERT INTO databank_cards (
                          subject_id, category, country, reporting_year,
                          title, content, source_type, source_url, reference_text, tags,
                          created_by, created_at, updated_at, is_active
                        )
                        SELECT %s, 'Glossary', NULL, NULL, %s, %s, 'System Seed', '', 'Seeded from default report glossary', 'report-glossary',
                               'system', NOW(), NOW(), TRUE
                        WHERE NOT EXISTS (
                          SELECT 1 FROM databank_cards
                          WHERE subject_id = %s AND lower(title) = lower(%s)
                        )
                        """,
                        [seed_subject_id, term, definition, seed_subject_id, term],
                    )
            _glossary_schema_seeded = True

        subject_row = con.execute(
            "SELECT subject_id FROM databank_subjects_lookup WHERE lower(name) = 'glossary' LIMIT 1"
        ).fetchone()
        if not subject_row:
            return []
        subject_id = int(subject_row[0])

        report_country = str(job_data.get("country") or "").strip()

        params: list[Any] = [subject_id]
        where_country = ""
        if report_country:
            where_country = "AND (c.country IS NULL OR trim(c.country) = '' OR lower(c.country) = lower(%s) OR lower(c.country) = 'global')"
            params.append(report_country)

        df = con.execute(
            f"""
            SELECT c.title, c.content
            FROM databank_cards c
            WHERE c.subject_id = %s
              AND c.is_active = TRUE
              {where_country}
            ORDER BY lower(c.title)
            """
            ,
            params,
        ).df()

    terms: list[dict[str, str]] = []
    if df is not None and not df.empty:
        for _, row in df.iterrows():
            term = str(row.get("title") or "").strip()
            definition = str(row.get("content") or "").strip()
            if term and definition:
                terms.append({"term": term, "definition": definition})
    return terms


def _safe_filename(text: str | None, ext: str) -> str:
    raw = str(text or "report").strip().replace(" ", "_")
    safe = "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in raw).strip("_")
    return f"{safe or 'report'}.{ext}"


def _format_template_value(value: object, variable_type: str) -> str:
    if value is None:
        return ""
    txt = str(value).strip()
    if not txt:
        return ""
    if variable_type == "boolean":
        norm = txt.lower()
        if norm in {"true", "1", "yes", "y", "on"}:
            return "Yes"
        if norm in {"false", "0", "no", "n", "off"}:
            return "No"
    return txt


def _docx_add_section_heading(doc: "DocxDocument", text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    return heading


def _docx_add_variable_sections(
    doc: "DocxDocument",
    template_variables: dict[str, object],
    template_var_meta: dict[str, dict[str, object]],
):
    """Render template variables grouped by section and ordered by display_order."""
    try:
        from docx.shared import Pt as DocxPt
    except Exception:  # pragma: no cover
        logger.debug("DOCX shared styling import failed; continuing without DocxPt styling", exc_info=True)
        DocxPt = None

    if not template_variables:
        return

    grouped: dict[str, list[tuple[str, str, int]]] = {}
    for key, raw_value in template_variables.items():
        meta = template_var_meta.get(key, {})
        variable_type = str(meta.get("type") or "text")
        label = str(meta.get("label") or key)
        section = str(meta.get("section") or "General")
        order = int(meta.get("order") or 0)

        value = _format_template_value(raw_value, variable_type)
        if value == "":
            continue

        grouped.setdefault(section, []).append((label, value, order))

    if not grouped:
        return

    _docx_add_section_heading(doc, "Template Variables", level=1)
    for section in sorted(grouped.keys()):
        _docx_add_section_heading(doc, section, level=2)
        rows = sorted(grouped[section], key=lambda item: (item[2], item[0].lower()))
        for label, value, _order in rows:
            p = doc.add_paragraph()
            run_label = p.add_run(f"{label}: ")
            run_label.bold = True
            if DocxPt is not None:
                run_label.font.size = DocxPt(10)
            run_value = p.add_run(value)
            if DocxPt is not None:
                run_value.font.size = DocxPt(10)


def _apply_docx_brand_styles(doc) -> None:
    """Apply NZI brand colours and typography to a python-docx Document object."""
    try:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # Brand palette
        BRAND_GREEN = RGBColor(0x14, 0x53, 0x2d)   # deep forest green
        BRAND_LIGHT = RGBColor(0x1F, 0x7A, 0x46)   # lighter green
        TEXT_MAIN   = RGBColor(0x0F, 0x17, 0x2A)   # near-black
        TEXT_MUTED  = RGBColor(0x47, 0x55, 0x69)   # slate

        # Title (Heading 0) â€” for the cover title
        title_style = doc.styles['Title']
        title_style.font.name = 'Aptos'
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = TEXT_MAIN
        title_style.paragraph_format.space_after = Pt(6)

        # Heading 1 â€” brand green, 16pt
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Aptos'
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = BRAND_GREEN
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(4)

        # Heading 2 â€” lighter green, 13pt
        h2 = doc.styles['Heading 2']
        h2.font.name = 'Aptos'
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = BRAND_LIGHT
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(3)

        # Normal body text
        normal = doc.styles['Normal']
        normal.font.name = 'Aptos'
        normal.font.size = Pt(10)
        normal.font.color.rgb = TEXT_MAIN
        normal.paragraph_format.space_after = Pt(4)

    except Exception:
        logger.warning("DOCX brand styling failed; continuing with unstyled document", exc_info=True)


def _docx_shade_table_header(table) -> None:
    """Apply NZI green shading and white text to the first row of a table."""
    try:
        from docx.shared import RGBColor, Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        header_row = table.rows[0]
        for cell in header_row.cells:
            # Set background fill
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '14532d')
            tcPr.append(shd)
            # Set text colour and weight
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
                    run.font.size = Pt(9)
                if not para.runs and para.text:
                    run = para.add_run(para.text)
                    para.clear()
                    run = para.add_run(para.text)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
    except Exception:
        logger.warning("DOCX table header shading failed; continuing without header styling", exc_info=True)


def _docx_add_cover_page(doc, job_data: dict, generation_date: str, template_name: str, report_metadata: dict) -> None:
    """Insert a branded cover page as the first content of the document."""
    try:
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        BRAND_GREEN = RGBColor(0x14, 0x53, 0x2d)
        TEXT_MUTED  = RGBColor(0x47, 0x55, 0x69)

        # Logo / brand name line
        brand_p = doc.add_paragraph()
        brand_p.paragraph_format.space_before = Pt(0)
        brand_run = brand_p.add_run("NZI INSIGHTS PRO")
        brand_run.font.name = 'Aptos'
        brand_run.font.size = Pt(8)
        brand_run.font.bold = True
        brand_run.font.color.rgb = BRAND_GREEN
        brand_run.font.all_caps = True

        # Report title
        client_name = str(job_data.get('client_name') or 'Client')
        reporting_year = str(job_data.get('reporting_year') or '')
        title_text = f"{client_name} Carbon Emissions Report"
        if reporting_year:
            title_text += f"\n{reporting_year}"
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(24)
        title_p.paragraph_format.space_after = Pt(6)
        title_run = title_p.add_run(title_text)
        title_run.font.name = 'Aptos'
        title_run.font.size = Pt(26)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        # Subtitle / period
        rm = report_metadata or {}
        period = str(rm.get('current_reporting_period_label') or
                     ((job_data.get('period_start') or '') + ' â€“ ' + (job_data.get('period_end') or '')))
        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(period.strip(' â€“'))
        sub_run.font.name = 'Aptos'
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = TEXT_MUTED
        sub_p.paragraph_format.space_after = Pt(18)

        # Divider rule
        hr_p = doc.add_paragraph('â”€' * 60)
        hr_p.runs[0].font.color.rgb = BRAND_GREEN
        hr_p.runs[0].font.size = Pt(8)
        hr_p.paragraph_format.space_after = Pt(12)

        # Metadata grid (2-column table, borderless)
        meta_entries = [
            ("Job Number",       str(job_data.get('job_number') or '-')),
            ("Reporting Year",   str(job_data.get('reporting_year') or '-')),
            ("Generated",        generation_date),
            ("Template",         template_name or '-'),
            ("Consultant",       str(rm.get('consultant_name') or '-')),
            ("Client Signee",    str(rm.get('client_signee_name') or '-')),
        ]
        meta_table = doc.add_table(rows=len(meta_entries), cols=2)
        meta_table.style = 'Table Grid'
        for i, (label, value) in enumerate(meta_entries):
            meta_table.cell(i, 0).text = label
            meta_table.cell(i, 1).text = value
            for cell in meta_table.rows[i].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Aptos'
                        run.font.size = Pt(10)
                    if cell == meta_table.rows[i].cells[0]:
                        for run in para.runs:
                            run.font.bold = True
                            run.font.color.rgb = TEXT_MUTED

        # Page break after cover
        doc.add_page_break()

    except Exception:
        logger.warning("DOCX cover creation failed; continuing with body content", exc_info=True)
        doc.add_page_break()


def _docx_embed_chart(doc, base64_str: str, title: str = "", width_cm: float = 14.0) -> None:
    """Decode a base64 PNG chart and embed it as an image in the document."""
    if not base64_str:
        return
    try:
        import base64 as _b64
        import io as _io
        from docx.shared import Cm

        img_bytes = _b64.b64decode(base64_str)
        img_stream = _io.BytesIO(img_bytes)
        if title:
            p = doc.add_paragraph(title)
            p.runs[0].bold = True if p.runs else None
        doc.add_picture(img_stream, width=Cm(width_cm))
    except Exception:
        logger.warning("Chart embed failed for report section '%s'; continuing without image", title, exc_info=True)


def _resolve_selected_template(job_id: int, payload: GenerateReportPayload | None):
    """Resolve template/version selection with assignment validation."""
    selected_template = None
    requested_template_id = payload.template_id if payload else None
    requested_version_id = payload.version_id if payload else None

    if requested_template_id is not None or requested_version_id is not None:
        if requested_template_id is None or requested_version_id is None:
            raise HTTPException(
                status_code=400,
                detail=(
                    "Both template_id and version_id are required when explicitly selecting a report template. "
                    "Assign a template version first, then try again."
                ),
            )

        selected_template = _get_template_selection(int(requested_template_id), int(requested_version_id))
        if not selected_template or selected_template.get("version_id") is None:
            raise HTTPException(status_code=404, detail="Selected report template/version not found")

        assigned_pair = _get_job_assignment_template_and_version(int(job_id))
        if assigned_pair is None:
            raise HTTPException(
                status_code=400,
                detail=(
                    "No report template/version is assigned to this job. "
                    "Assign a template version in the Report Generator before generating."
                ),
            )

        assigned_template_id, assigned_version_id = assigned_pair
        if int(requested_template_id) != assigned_template_id or int(requested_version_id) != assigned_version_id:
            raise HTTPException(
                status_code=400,
                detail=(
                    "The requested template/version does not match this jobâ€™s assigned template version. "
                    "Re-assign the template version and try again."
                ),
            )

        return selected_template

    assigned_pair = _get_job_assignment_template_and_version(int(job_id))
    if assigned_pair is None:
        raise HTTPException(
            status_code=400,
            detail=(
                "No report template/version is assigned to this job. "
                "Please assign a report template version before generating."
            ),
        )

    assigned_template_id, assigned_version_id = assigned_pair
    selected_template = _get_template_selection(int(assigned_template_id), int(assigned_version_id))
    if not selected_template or selected_template.get("version_id") is None:
        raise HTTPException(
            status_code=404,
            detail="The assigned report template/version could not be loaded. Please reassign and try again.",
        )

    return selected_template


def _stringify_render_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Yes" if value else "No"
    return str(value)


def _apply_braced_placeholder_substitution(content: str, render_values: dict[str, object]) -> str:
    """
    Replace one-to-one legacy placeholders like {Report Title} using merged render values.
    Leaves unknown tokens untouched.
    """
    if not content:
        return content

    def _replace(match: re.Match[str]) -> str:
        token = match.group(0)
        bare = token[1:-1].strip()

        if token in render_values:
            return _stringify_render_value(render_values.get(token))
        if bare in render_values:
            return _stringify_render_value(render_values.get(bare))
        return token

    # Ignore Jinja-style double braces (e.g. {{ job_data.client_name }})
    return re.sub(r"(?<!\{)\{[^{}]+\}(?!\})", _replace, content)


def get_emissions_by_activity(job_id: int, _user: dict[str, str] = Depends(_current_user)):
    """
    Get detailed emissions breakdown by activity with category groupings.
    Returns data for donut chart, bar chart, and detailed table.
    """
    try:
        categories = get_emissions_by_category(job_id)
        activity_groups, activity_totals, activity_details = _build_activity_grouping(categories)
    except Exception:
        logger.warning("Failed to build activity grouping for job %s; using empty categories", job_id, exc_info=True)
        categories = []
        activity_groups, activity_totals, activity_details = _build_activity_grouping(categories)

    return {
        'activity_groups': activity_groups,
        'activity_totals': activity_totals,
        'activity_details': activity_details,
        'activity_group_order': ACTIVITY_GROUP_ORDER,
        'activity_group_colors': ACTIVITY_GROUP_COLORS,
        'all_categories': categories,
    }


def get_report_assets(job_id: int, _user: dict[str, str] = Depends(_current_user)):
    """
    Generate all required chart assets for a carbon report.
    This endpoint generates all charts and validates they're present before PDF generation.
    
    Returns:
        Dictionary of chart assets (base64 encoded PNGs)
        Includes validation status and any missing asset warnings
    """
    try:
        # Fetch job data
        job_data = get_job_data(job_id)
        if not job_data:
            raise HTTPException(status_code=404, detail="Job not found")

        if not save_version:
            with get_conn() as con:
                frozen_snapshot = _load_latest_final_report_version_snapshot(con, int(job_id))
            if frozen_snapshot:
                version_row, snapshot_payload = frozen_snapshot
                html_content = _render_report_snapshot_html(snapshot_payload)
                pdf_bytes = _render_html_to_pdf_bytes(html_content)
                headers = {
                    "Content-Disposition": f"inline; filename=job-{job_id}-emissions-report.pdf",
                    "X-Report-Version-Saved": "false",
                    "X-Report-Snapshot-Hash": str(version_row.get("data_hash") or ""),
                    "X-Report-Version-Id": str(version_row.get("report_version_id") or ""),
                    "X-Report-Version-Number": str(version_row.get("version_number") or ""),
                    "X-Report-Version-Label": str(version_row.get("version_label") or ""),
                    "X-Report-Version-Status": str(version_row.get("status") or ""),
                    "X-Report-File-Id": str(version_row.get("file_id") or ""),
                }
                return Response(
                    content=pdf_bytes,
                    media_type="application/pdf",
                    headers=headers,
                )
        
        scope_totals = get_scope_totals(job_id)
        categories = get_emissions_by_category(job_id)
        activity_groups, activity_totals, _details = _build_activity_grouping(categories)
        
        # Get intensity metrics (optional)
        intensity_metrics = get_intensity_metrics(job_id)
        
        # Get target data (for reduction pathway charts)
        target_data = get_job_target_data(job_id)
        
        # Generate all chart assets
        assets = generate_report_assets(
            job_data=job_data,
            scope_totals=scope_totals,
            categories=categories,
            activity_totals=activity_totals,
            intensity_metrics=intensity_metrics,
            target_data=target_data
        )
        
        # Validate required assets
        is_valid, missing_keys = validate_report_assets(assets)
        
        return {
            'job_id': job_id,
            'assets': assets,
            'validation': {
                'is_valid': is_valid,
                'missing_keys': missing_keys,
                'required_assets': REQUIRED_ASSETS,
                'generated_assets': list(assets.keys())
            },
            'target_data': target_data,
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate report assets: {str(e)}")


def generate_report_with_assets(
    job_id: int,
    request: Request,
    skip_validation: bool = Query(False),
    save_version: bool = Query(False),
    report_version_status: str = Query("review"),
    report_version_label: str | None = Query(None),
    report_version_notes: str | None = Query(None),
    _user: dict[str, str] = Depends(_current_user),
):
    """
    Generate a comprehensive PDF report with all required chart assets.
    Includes validation step to ensure all required assets are present before PDF generation.
    
    Args:
        job_id: The job ID to generate the report for
        skip_validation: If true, skip the assets validation check (not recommended)
    """
    try:
        # Fetch job data
        job_data = get_job_data(job_id)
        if not job_data:
            raise HTTPException(status_code=404, detail="Job not found")

        if not save_version:
            with get_conn() as con:
                frozen_snapshot = _load_latest_final_report_version_snapshot(con, int(job_id))
                if frozen_snapshot:
                    version_row, snapshot_payload = frozen_snapshot
                    file_payload = None
                    file_id = version_row.get("file_id")
                    if file_id is not None:
                        file_payload = _load_report_version_file_payload(con, int(job_id), int(file_id))
                    if file_payload:
                        headers = {
                            "Content-Disposition": f"inline; filename={file_payload['file_name']}",
                            "X-Report-Version-Saved": "false",
                            "X-Report-Snapshot-Hash": str(version_row.get("data_hash") or ""),
                            "X-Report-Version-Id": str(version_row.get("report_version_id") or ""),
                            "X-Report-Version-Number": str(version_row.get("version_number") or ""),
                            "X-Report-Version-Label": str(version_row.get("version_label") or ""),
                            "X-Report-Version-Status": str(version_row.get("status") or ""),
                            "X-Report-File-Id": str(version_row.get("file_id") or ""),
                        }
                        return Response(
                            content=file_payload["content"],
                            media_type=str(file_payload.get("media_type") or "application/octet-stream"),
                            headers=headers,
                        )
                    html_content = _render_report_snapshot_html(snapshot_payload)
                    pdf_bytes = _render_html_to_pdf_bytes(html_content)
                    headers = {
                        "Content-Disposition": f"inline; filename=job-{job_id}-emissions-report.pdf",
                        "X-Report-Version-Saved": "false",
                        "X-Report-Snapshot-Hash": str(version_row.get("data_hash") or ""),
                        "X-Report-Version-Id": str(version_row.get("report_version_id") or ""),
                        "X-Report-Version-Number": str(version_row.get("version_number") or ""),
                        "X-Report-Version-Label": str(version_row.get("version_label") or ""),
                        "X-Report-Version-Status": str(version_row.get("status") or ""),
                        "X-Report-File-Id": str(version_row.get("file_id") or ""),
                    }
                    return Response(
                        content=pdf_bytes,
                        media_type="application/pdf",
                        headers=headers,
                    )
        
        scope_totals = get_scope_totals(job_id)
        categories = get_emissions_by_category(job_id)
        benchmark_totals = get_benchmark_emissions(job_id, job_data.get('benchmark_year'))
        site_breakdowns = get_site_emissions_breakdowns(job_id)
        
        generation_date = datetime.now().strftime('%d %B %Y')

        # Get activity grouping
        activity_groups, activity_totals, activity_details = _build_activity_grouping(categories)
        benchmark_job_id = _get_benchmark_job_id(job_id, job_data.get("benchmark_year"))
        benchmark_categories = get_emissions_by_category(benchmark_job_id) if benchmark_job_id else []
        _bg, benchmark_activity_totals, _bd = _build_activity_grouping(benchmark_categories) if benchmark_categories else ({}, {}, {})
        benchmark_site_breakdowns = get_site_emissions_breakdowns(benchmark_job_id) if benchmark_job_id else {
            "show_site_tables": False,
            "site_count": 0,
            "overall": [],
            "scope": [],
            "activity": [],
            "appendix_rows": [],
        }
        benchmark_job_id = _get_benchmark_job_id(job_id, job_data.get("benchmark_year"))
        benchmark_categories = get_emissions_by_category(benchmark_job_id) if benchmark_job_id else []
        _bg, benchmark_activity_totals, _bd = _build_activity_grouping(benchmark_categories) if benchmark_categories else ({}, {}, {})
        benchmark_site_breakdowns = get_site_emissions_breakdowns(benchmark_job_id) if benchmark_job_id else {
            "show_site_tables": False,
            "site_count": 0,
            "overall": [],
            "scope": [],
            "activity": [],
            "appendix_rows": [],
        }
        
        # Get intensity metrics
        intensity_metrics = get_intensity_metrics(job_id)
        
        # Get target data
        target_data = get_job_target_data(job_id)
        glossary_terms = _ensure_glossary_cards_and_fetch(job_id, job_data, target_data.get('glossary_terms'))
        if not glossary_terms:
            glossary_terms = list(DEFAULT_GLOSSARY_ENTRIES)
        
        # Generate all chart assets using the new comprehensive function
        assets = generate_report_assets(
            job_data=job_data,
            scope_totals=scope_totals,
            categories=categories,
            activity_totals=activity_totals,
            intensity_metrics=intensity_metrics,
            target_data=target_data
        )
        
        # Validate required assets (unless skipped)
        if not skip_validation:
            is_valid, missing_keys = validate_report_assets(assets)
            if not is_valid:
                raise HTTPException(
                    status_code=400,
                    detail={
                        'error': 'Missing required chart assets',
                        'missing_keys': missing_keys,
                        'required_assets': REQUIRED_ASSETS,
                        'hint': 'Use skip_validation=true to bypass this check (not recommended for production)'
                    }
                )
        
        # Resolve template selection
        selected_template = _get_job_assigned_template_selection(int(job_id))
        template_variables = {}
        render_meta = None
        
        if selected_template and selected_template.get("version_id") is not None:
            template_variables = _get_template_variable_values_for_render(
                int(job_id),
                int(selected_template["template_id"]),
                int(selected_template["version_id"]),
            )
            render_meta = {
                "template_id": selected_template.get("template_id"),
                "template_key": selected_template.get("template_key"),
                "template_name": selected_template.get("template_name"),
                "template_type": selected_template.get("template_type"),
                "report_type": selected_template.get("report_type"),
                "version_id": selected_template.get("version_id"),
                "version_number": selected_template.get("version_number"),
                "version_label": selected_template.get("version_label"),
                "version_status": selected_template.get("version_status"),
                "source": "assignment",
            }
        
        is_yoy_template = str((selected_template or {}).get("template_key") or "").strip().lower() == "crp_yoy_benchmark"
        scope_comparison = _build_scope_comparison(scope_totals, benchmark_totals)
        activity_comparison = _build_activity_comparison(activity_totals, benchmark_activity_totals)
        site_overall_comparison = _build_site_overall_comparison(site_breakdowns, benchmark_site_breakdowns)

        report_metadata = _get_job_report_metadata(
            int(job_id),
            updated_by=_user.get("email", "unknown"),
        )
        job_actions = get_job_report_actions_payload(int(job_id))
        render_values = _build_report_render_values(
            job_data=job_data,
            scope_totals=scope_totals,
            template_variables=template_variables,
            report_metadata=report_metadata,
            generation_date=generation_date,
        )
        
        # Setup Jinja2 template environment
        template_dir = os.path.join(os.path.dirname(__file__), 'templates')
        env = Environment(loader=FileSystemLoader(template_dir))
        
        # Use interactive report template which supports assets
        template = env.get_template('interactive_report.html')
        
        # Prepare activity data for template
        activity_labels = [g for g in ACTIVITY_GROUP_ORDER if float(activity_totals.get(g, 0) or 0) > 0]
        if not activity_labels:
            activity_labels = ACTIVITY_GROUP_ORDER
        activity_values = [float(activity_totals.get(g, 0) or 0) for g in activity_labels]
        
        # Render HTML with data and all assets
        html_content = template.render(
            job_data=job_data,
            scope_totals=scope_totals,
            benchmark_totals=benchmark_totals,
            categories=categories,
            activity_groups=activity_groups,
            activity_totals=activity_totals,
            activity_details=activity_details,
            activity_group_order=ACTIVITY_GROUP_ORDER,
            activity_group_colors=ACTIVITY_GROUP_COLORS,
            activity_labels=activity_labels,
            activity_values=activity_values,
            # Pass assets dictionary for template to use
            assets=assets,
            # Legacy individual chart support (for backward compatibility)
            scope_chart_base64=assets.get('scope_breakdown', ''),
            activity_chart_base64=assets.get('activity_breakdown', ''),
            generation_date=generation_date,
            template_variables=template_variables,
            report_metadata=report_metadata,
            job_actions=job_actions,
            nzi_logo_src=_get_nzi_logo_src(),
            client_logo_src=_get_client_logo_src(job_data.get("logo_url")),
            render_values=render_values,
            render_template=render_meta,
            glossary_terms=glossary_terms,
            # Pass target data for reduction pathway charts
            target_data=target_data,
            site_breakdowns=site_breakdowns,
            is_yoy_template=is_yoy_template,
            scope_comparison=scope_comparison,
            benchmark_activity_totals=benchmark_activity_totals,
            activity_comparison=activity_comparison,
            benchmark_site_breakdowns=benchmark_site_breakdowns,
            site_overall_comparison=site_overall_comparison,
        )

        html_content = _apply_braced_placeholder_substitution(html_content, render_values)

        report_snapshot_payload, report_snapshot_hash = _build_report_version_snapshot(
            job_data=job_data,
            selected_template=render_meta,
            template_variables=template_variables,
            report_metadata=report_metadata,
            job_actions=job_actions,
            scope_totals=scope_totals,
            benchmark_totals=benchmark_totals,
            categories=categories,
            activity_totals=activity_totals,
            site_breakdowns=site_breakdowns,
            benchmark_site_breakdowns=benchmark_site_breakdowns,
            render_values=render_values,
            target_data=target_data,
            scope_comparison=scope_comparison,
            activity_comparison=activity_comparison,
            site_overall_comparison=site_overall_comparison,
            assets=assets,
            generation_date=generation_date,
        )
        report_snapshot_json = json.dumps(report_snapshot_payload, ensure_ascii=False, sort_keys=True, default=str)
        normalized_version_status = _normalize_report_version_status(report_version_status)
        
        pdf_bytes = _render_html_to_pdf_bytes(html_content)

        saved_report_version: dict[str, Any] | None = None
        if save_version:
            with get_conn(autocommit=False) as con:
                _ensure_job_files_table(con)
                _ensure_report_versions_schema(con)
                version_row = con.execute(
                    "SELECT COALESCE(MAX(version_number), 0) + 1 FROM job_report_versions WHERE job_id = %s",
                    [int(job_id)],
                ).fetchone()
                next_version_number = int(version_row[0]) if version_row and version_row[0] is not None else 1
                saved_report_version = _store_report_version_artifact(
                    con=con,
                    job_data=job_data,
                    version_number=next_version_number,
                    pdf_bytes=pdf_bytes,
                    generated_by=_user.get("email", "unknown"),
                    status=normalized_version_status,
                    template_id=int(selected_template["template_id"]) if selected_template and selected_template.get("template_id") is not None else None,
                    template_version_id=int(selected_template["version_id"]) if selected_template and selected_template.get("version_id") is not None else None,
                    data_hash=report_snapshot_hash,
                    snapshot_json=report_snapshot_json,
                    version_label=report_version_label,
                    notes=report_version_notes,
                )
                record_audit_event(
                    con,
                    request=request,
                    actor=_user,
                    action="create",
                    entity_type="job_report_version",
                    entity_id=saved_report_version.get("report_version_id"),
                    client_id=int(job_data.get("client_db_id") or 0) if job_data.get("client_db_id") is not None else None,
                    job_id=int(job_id),
                    after={
                        "report_version_id": saved_report_version.get("report_version_id"),
                        "version_number": saved_report_version.get("version_number"),
                        "version_label": saved_report_version.get("version_label"),
                        "status": saved_report_version.get("status"),
                        "file_id": saved_report_version.get("file_id"),
                        "data_hash": report_snapshot_hash,
                    },
                    metadata={
                        "report_snapshot_hash": report_snapshot_hash,
                        "report_version_status": normalized_version_status,
                        "save_version": True,
                    },
                )
        
        # Return PDF
        headers = {
            "Content-Disposition": f"inline; filename=job-{job_id}-emissions-report.pdf",
            "X-Report-Version-Saved": "true" if saved_report_version else "false",
            "X-Report-Snapshot-Hash": report_snapshot_hash,
        }
        if saved_report_version:
            headers.update(
                {
                    "X-Report-Version-Id": str(saved_report_version.get("report_version_id") or ""),
                    "X-Report-Version-Number": str(saved_report_version.get("version_number") or ""),
                    "X-Report-Version-Label": str(saved_report_version.get("version_label") or ""),
                    "X-Report-Version-Status": str(saved_report_version.get("status") or ""),
                    "X-Report-File-Id": str(saved_report_version.get("file_id") or ""),
                }
            )

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers=headers,
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate report: {str(e)}")


def generate_job_report(
    job_id: int,
    request: Request,
    payload: GenerateReportPayload | None = None,
    _user: dict[str, str] = Depends(_current_user),
):
    """
    Generate a comprehensive PDF emissions report for a job using HTML templates.
    """
    try:
        # Fetch data
        job_data = get_job_data(job_id)
        if not job_data:
            raise HTTPException(status_code=404, detail="Job not found")

        save_version = bool(payload.save_version) if payload else False
        if not save_version:
            with get_conn() as con:
                frozen_snapshot = _load_latest_final_report_version_snapshot(con, int(job_id))
                if frozen_snapshot:
                    version_row, snapshot_payload = frozen_snapshot
                    file_payload = None
                    file_id = version_row.get("file_id")
                    if file_id is not None:
                        file_payload = _load_report_version_file_payload(con, int(job_id), int(file_id))
                    if file_payload:
                        headers = {
                            "Content-Disposition": f"inline; filename={file_payload['file_name']}",
                            "X-Report-Version-Saved": "false",
                            "X-Report-Snapshot-Hash": str(version_row.get("data_hash") or ""),
                            "X-Report-Version-Id": str(version_row.get("report_version_id") or ""),
                            "X-Report-Version-Number": str(version_row.get("version_number") or ""),
                            "X-Report-Version-Label": str(version_row.get("version_label") or ""),
                            "X-Report-Version-Status": str(version_row.get("status") or ""),
                            "X-Report-File-Id": str(version_row.get("file_id") or ""),
                        }
                        return Response(
                            content=file_payload["content"],
                            media_type=str(file_payload.get("media_type") or "application/octet-stream"),
                            headers=headers,
                        )
                    html_content = _render_report_snapshot_html(snapshot_payload)
                    pdf_bytes = _render_html_to_pdf_bytes(html_content)
                    headers = {
                        "Content-Disposition": f"inline; filename=job-{job_id}-emissions-report.pdf",
                        "X-Report-Version-Saved": "false",
                        "X-Report-Snapshot-Hash": str(version_row.get("data_hash") or ""),
                        "X-Report-Version-Id": str(version_row.get("report_version_id") or ""),
                        "X-Report-Version-Number": str(version_row.get("version_number") or ""),
                        "X-Report-Version-Label": str(version_row.get("version_label") or ""),
                        "X-Report-Version-Status": str(version_row.get("status") or ""),
                        "X-Report-File-Id": str(version_row.get("file_id") or ""),
                    }
                    return Response(
                        content=pdf_bytes,
                        media_type="application/pdf",
                        headers=headers,
                    )
        
        scope_totals = get_scope_totals(job_id)
        categories = get_emissions_by_category(job_id)
        benchmark_totals = get_benchmark_emissions(job_id, job_data.get('benchmark_year'))
        site_breakdowns = get_site_emissions_breakdowns(job_id)
        
        generation_date = datetime.now().strftime('%d %B %Y')

        # Generate charts as base64 images
        period_text = f"{job_data.get('period_start', 'N/A')} to {job_data.get('period_end', 'N/A')}"
        
        # Scope chart
        scope_chart_data = {
            'Scope 1': scope_totals['Scope 1'],
            'Scope 2': scope_totals['Scope 2'],
            'Scope 3': scope_totals['Scope 3']
        }
        scope_colors = {
            'Scope 1': '#4472C4',
            'Scope 2': '#ED7D31',
            'Scope 3': '#70AD47'
        }
        scope_chart_base64 = create_donut_chart(
            scope_chart_data,
            scope_colors,
            f'Emissions by Scope (tCOâ‚‚e)\n{period_text}',
            scope_totals['Total']
        )
        
        # Activity chart + detailed grouped rows
        activity_groups, activity_totals, activity_details = _build_activity_grouping(categories)
        benchmark_job_id = _get_benchmark_job_id(job_id, job_data.get("benchmark_year"))
        benchmark_categories = get_emissions_by_category(benchmark_job_id) if benchmark_job_id else []
        _bg, benchmark_activity_totals, _bd = _build_activity_grouping(benchmark_categories) if benchmark_categories else ({}, {}, {})
        benchmark_site_breakdowns = get_site_emissions_breakdowns(benchmark_job_id) if benchmark_job_id else {
            "show_site_tables": False,
            "site_count": 0,
            "overall": [],
            "scope": [],
            "activity": [],
            "appendix_rows": [],
        }
        period_from = job_data.get("period_start")
        period_to = job_data.get("period_end")
        period_from = job_data.get("period_start")
        period_to = job_data.get("period_end")

        # Activity chart - using new professional donut style
        activity_chart_base64 = create_activity_donut(
            activity_totals,
            ACTIVITY_GROUP_COLORS,
            title="Emissions by Activity",
            center_label="tCO₂e",
            period_from=period_from,
            period_to=period_to
        )
        
        # Resolve template selection priority:
        # 1) Explicit payload template/version (both required)
        # 2) Job assignment template/version (both required)
        # 3) No fallback for this endpoint; fail fast with clear guidance
        selected_template = None
        requested_template_id = payload.template_id if payload else None
        requested_version_id = payload.version_id if payload else None

        if requested_template_id is not None or requested_version_id is not None:
            if requested_template_id is None or requested_version_id is None:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "Both template_id and version_id are required when explicitly selecting a report template. "
                        "Assign a template version first, then try again."
                    ),
                )
            selected_template = _get_template_selection(
                int(requested_template_id),
                int(requested_version_id),
            )
            if not selected_template or selected_template.get("version_id") is None:
                raise HTTPException(status_code=404, detail="Selected report template/version not found")

            # Ensure requested template/version matches assignment for this job
            assigned_pair = _get_job_assignment_template_and_version(int(job_id))
            if assigned_pair is None:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "No report template/version is assigned to this job. "
                        "Assign a template version in the Report Generator before generating."
                    ),
                )

            assigned_template_id, assigned_version_id = assigned_pair
            if int(requested_template_id) != assigned_template_id or int(requested_version_id) != assigned_version_id:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "The requested template/version does not match this jobâ€™s assigned template version. "
                        "Re-assign the template version and try again."
                    ),
                )
        else:
            assigned_pair = _get_job_assignment_template_and_version(int(job_id))
            if assigned_pair is None:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "No report template/version is assigned to this job. "
                        "Please assign a report template version before generating."
                    ),
                )
            assigned_template_id, assigned_version_id = assigned_pair
            selected_template = _get_template_selection(int(assigned_template_id), int(assigned_version_id))
            if not selected_template or selected_template.get("version_id") is None:
                raise HTTPException(
                    status_code=404,
                    detail="The assigned report template/version could not be loaded. Please reassign and try again.",
                )

        template_variables = {}
        if selected_template:
            template_variables = _get_template_variable_values_for_render(
                int(job_id),
                int(selected_template["template_id"]),
                int(selected_template["version_id"]) if selected_template.get("version_id") is not None else None,
            )

        report_metadata = _get_job_report_metadata(
            int(job_id),
            updated_by=_user.get("email", "unknown"),
        )
        job_actions = get_job_report_actions_payload(int(job_id))

        render_values = _build_report_render_values(
            job_data=job_data,
            scope_totals=scope_totals,
            template_variables=template_variables,
            report_metadata=report_metadata,
            generation_date=generation_date,
        )

        # Setup Jinja2 template environment
        template_dir = os.path.join(os.path.dirname(__file__), 'templates')
        env = Environment(loader=FileSystemLoader(template_dir))

        template = _resolve_template_for_render(env, selected_template)

        render_meta = {
            "template_id": selected_template.get("template_id") if selected_template else None,
            "template_key": selected_template.get("template_key") if selected_template else None,
            "template_name": selected_template.get("template_name") if selected_template else None,
            "template_type": selected_template.get("template_type") if selected_template else None,
            "report_type": selected_template.get("report_type") if selected_template else None,
            "version_id": selected_template.get("version_id") if selected_template else None,
            "version_number": selected_template.get("version_number") if selected_template else None,
            "version_label": selected_template.get("version_label") if selected_template else None,
            "version_status": selected_template.get("version_status") if selected_template else None,
            "source": "assignment_or_request" if selected_template else "unreachable",
        }
        is_yoy_template = str((selected_template or {}).get("template_key") or "").strip().lower() == "crp_yoy_benchmark"
        scope_comparison = _build_scope_comparison(scope_totals, benchmark_totals)
        activity_comparison = _build_activity_comparison(activity_totals, benchmark_activity_totals)
        site_overall_comparison = _build_site_overall_comparison(site_breakdowns, benchmark_site_breakdowns)
        
        # Get target data
        target_data = get_job_target_data(job_id)
        glossary_terms = _ensure_glossary_cards_and_fetch(job_id, job_data, target_data.get('glossary_terms'))
        if not glossary_terms:
            glossary_terms = list(DEFAULT_GLOSSARY_ENTRIES)

        # Get intensity metrics for chart generation
        intensity_metrics = get_intensity_metrics(job_id)
        
        # Generate chart assets
        assets = generate_report_assets(
            job_data=job_data,
            scope_totals=scope_totals,
            categories=categories,
            activity_totals=activity_totals,
            intensity_metrics=intensity_metrics,
            target_data=target_data
        )
        
        # Render HTML with data
        html_content = template.render(
            job_data=job_data,
            scope_totals=scope_totals,
            benchmark_totals=benchmark_totals,
            categories=categories,
            activity_groups=activity_groups,
            activity_totals=activity_totals,
            activity_details=activity_details,
            activity_group_order=ACTIVITY_GROUP_ORDER,
            activity_group_colors=ACTIVITY_GROUP_COLORS,
            assets=assets,
            scope_chart_base64=scope_chart_base64,
            activity_chart_base64=activity_chart_base64,
            generation_date=generation_date,
            template_variables=template_variables,
            report_metadata=report_metadata,
            job_actions=job_actions,
            nzi_logo_src=_get_nzi_logo_src(),
            client_logo_src=_get_client_logo_src(job_data.get("logo_url")),
            render_values=render_values,
            render_template=render_meta,
            glossary_terms=glossary_terms,
            target_data=target_data,
            site_breakdowns=site_breakdowns,
            is_yoy_template=is_yoy_template,
            scope_comparison=scope_comparison,
            benchmark_activity_totals=benchmark_activity_totals,
            activity_comparison=activity_comparison,
            benchmark_site_breakdowns=benchmark_site_breakdowns,
            site_overall_comparison=site_overall_comparison,
        )

        html_content = _apply_braced_placeholder_substitution(html_content, render_values)

        report_snapshot_payload, report_snapshot_hash = _build_report_version_snapshot(
            job_data=job_data,
            selected_template=render_meta,
            template_variables=template_variables,
            report_metadata=report_metadata,
            job_actions=job_actions,
            scope_totals=scope_totals,
            benchmark_totals=benchmark_totals,
            categories=categories,
            activity_totals=activity_totals,
            site_breakdowns=site_breakdowns,
            benchmark_site_breakdowns=benchmark_site_breakdowns,
            render_values=render_values,
            target_data=target_data,
            scope_comparison=scope_comparison,
            activity_comparison=activity_comparison,
            site_overall_comparison=site_overall_comparison,
            assets=assets,
            generation_date=generation_date,
        )
        report_snapshot_json = json.dumps(report_snapshot_payload, ensure_ascii=False, sort_keys=True, default=str)
        normalized_version_status = _normalize_report_version_status(
            getattr(payload, "report_version_status", None) if payload else None
        )
        
        pdf_bytes = _render_html_to_pdf_bytes(html_content)

        saved_report_version: dict[str, Any] | None = None
        should_save_version = bool(payload.save_version) if payload else False
        if should_save_version:
            with get_conn(autocommit=False) as con:
                _ensure_job_files_table(con)
                _ensure_report_versions_schema(con)
                version_row = con.execute(
                    "SELECT COALESCE(MAX(version_number), 0) + 1 FROM job_report_versions WHERE job_id = %s",
                    [int(job_id)],
                ).fetchone()
                next_version_number = int(version_row[0]) if version_row and version_row[0] is not None else 1
                saved_report_version = _store_report_version_artifact(
                    con=con,
                    job_data=job_data,
                    version_number=next_version_number,
                    pdf_bytes=pdf_bytes,
                    generated_by=_user.get("email", "unknown"),
                    status=normalized_version_status,
                    template_id=int(selected_template["template_id"]) if selected_template and selected_template.get("template_id") is not None else None,
                    template_version_id=int(selected_template["version_id"]) if selected_template and selected_template.get("version_id") is not None else None,
                    data_hash=report_snapshot_hash,
                    snapshot_json=report_snapshot_json,
                    version_label=getattr(payload, "report_version_label", None) if payload else None,
                    notes=getattr(payload, "report_version_notes", None) if payload else None,
                )
                record_audit_event(
                    con,
                    request=request,
                    actor=_user,
                    action="create",
                    entity_type="job_report_version",
                    entity_id=saved_report_version.get("report_version_id"),
                    client_id=int(job_data.get("client_db_id") or 0) if job_data.get("client_db_id") is not None else None,
                    job_id=int(job_id),
                    after={
                        "report_version_id": saved_report_version.get("report_version_id"),
                        "version_number": saved_report_version.get("version_number"),
                        "version_label": saved_report_version.get("version_label"),
                        "status": saved_report_version.get("status"),
                        "file_id": saved_report_version.get("file_id"),
                        "data_hash": report_snapshot_hash,
                    },
                    metadata={
                        "report_snapshot_hash": report_snapshot_hash,
                        "report_version_status": normalized_version_status,
                        "save_version": True,
                    },
                )

        # Return PDF
        headers = {
            "Content-Disposition": f"inline; filename=job-{job_id}-emissions-report.pdf",
            "X-Report-Version-Saved": "true" if saved_report_version else "false",
            "X-Report-Snapshot-Hash": report_snapshot_hash,
        }
        if saved_report_version:
            headers.update(
                {
                    "X-Report-Version-Id": str(saved_report_version.get("report_version_id") or ""),
                    "X-Report-Version-Number": str(saved_report_version.get("version_number") or ""),
                    "X-Report-Version-Label": str(saved_report_version.get("version_label") or ""),
                    "X-Report-Version-Status": str(saved_report_version.get("status") or ""),
                    "X-Report-File-Id": str(saved_report_version.get("file_id") or ""),
                }
            )

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers=headers,
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate report: {str(e)}")


def generate_html_report(
    job_id: int,
    template_id: int | None = Query(None),
    version_id: int | None = Query(None),
    _user: dict[str, str] = Depends(_current_user),
):
    """
    Generate an interactive HTML report that can be viewed in browser or printed to PDF.
    This is the modern, dual-purpose report format.
    """
    try:
        # Fetch data
        job_data = get_job_data(job_id)
        if not job_data:
            raise HTTPException(status_code=404, detail="Job not found")

        with get_conn() as con:
            frozen_snapshot = _load_latest_final_report_version_snapshot(con, int(job_id))
        if frozen_snapshot:
            version_row, snapshot_payload = frozen_snapshot
            html_content = _render_report_snapshot_html(snapshot_payload)
            label = version_row.get("version_label") or f"v{version_row.get('version_number') or job_id}"
            return HTMLResponse(
                content=html_content,
                headers={
                    "X-Report-Version-Id": str(version_row.get("report_version_id") or ""),
                    "X-Report-Version-Label": str(label),
                    "X-Report-Version-Status": str(version_row.get("status") or ""),
                    "Cache-Control": "no-store",
                },
            )
        
        scope_totals = get_scope_totals(job_id)
        categories = get_emissions_by_category(job_id)
        benchmark_totals = get_benchmark_emissions(job_id, job_data.get('benchmark_year'))
        site_breakdowns = get_site_emissions_breakdowns(job_id)
        
        activity_groups, activity_totals, activity_details = _build_activity_grouping(categories)
        benchmark_job_id = _get_benchmark_job_id(job_id, job_data.get("benchmark_year"))
        benchmark_categories = get_emissions_by_category(benchmark_job_id) if benchmark_job_id else []
        _bg, benchmark_activity_totals, _bd = _build_activity_grouping(benchmark_categories) if benchmark_categories else ({}, {}, {})
        benchmark_site_breakdowns = get_site_emissions_breakdowns(benchmark_job_id) if benchmark_job_id else {
            "show_site_tables": False,
            "site_count": 0,
            "overall": [],
            "scope": [],
            "activity": [],
            "appendix_rows": [],
        }

        # Prepare data for charts (ordered + non-zero by default)
        activity_labels = [g for g in ACTIVITY_GROUP_ORDER if float(activity_totals.get(g, 0) or 0) > 0]
        if not activity_labels:
            activity_labels = ACTIVITY_GROUP_ORDER
        activity_values = [float(activity_totals.get(g, 0) or 0) for g in activity_labels]

        selected_template = _resolve_selected_template(
            int(job_id),
            GenerateReportPayload(template_id=template_id, version_id=version_id),
        )
        forced_builtin_template = _use_forced_builtin_template(selected_template)
        if forced_builtin_template:
            template_source = "builtin"
        elif selected_template and str(selected_template.get("template_content") or "").strip():
            template_source = "database"
        else:
            template_source = "builtin-fallback"
        template_variables = {}
        render_meta = None
        if selected_template and selected_template.get("version_id") is not None:
            template_variables = _get_template_variable_values_for_render(
                int(job_id),
                int(selected_template["template_id"]),
                int(selected_template["version_id"]),
            )
            render_meta = {
                "template_id": selected_template.get("template_id"),
                "template_key": selected_template.get("template_key"),
                "template_name": selected_template.get("template_name"),
                "template_type": selected_template.get("template_type"),
                "report_type": selected_template.get("report_type"),
                "version_id": selected_template.get("version_id"),
                "version_number": selected_template.get("version_number"),
                "version_label": selected_template.get("version_label"),
                "version_status": selected_template.get("version_status"),
                "source": "assignment",
            }
        is_yoy_template = str((selected_template or {}).get("template_key") or "").strip().lower() == "crp_yoy_benchmark"
        scope_comparison = _build_scope_comparison(scope_totals, benchmark_totals)
        activity_comparison = _build_activity_comparison(activity_totals, benchmark_activity_totals)
        site_overall_comparison = _build_site_overall_comparison(site_breakdowns, benchmark_site_breakdowns)
        is_yoy_template = str((selected_template or {}).get("template_key") or "").strip().lower() == "crp_yoy_benchmark"
        scope_comparison = _build_scope_comparison(scope_totals, benchmark_totals)
        activity_comparison = _build_activity_comparison(activity_totals, benchmark_activity_totals)
        site_overall_comparison = _build_site_overall_comparison(site_breakdowns, benchmark_site_breakdowns)

        generation_date = datetime.now().strftime('%d %B %Y')
        report_metadata = _get_job_report_metadata(
            int(job_id),
            updated_by=_user.get("email", "unknown"),
        )
        job_actions = get_job_report_actions_payload(int(job_id))
        render_values = _build_report_render_values(
            job_data=job_data,
            scope_totals=scope_totals,
            template_variables=template_variables,
            report_metadata=report_metadata,
            generation_date=generation_date,
        )
        
        # Get target data for the template
        target_data = get_job_target_data(job_id)
        glossary_terms = _ensure_glossary_cards_and_fetch(job_id, job_data, target_data.get('glossary_terms'))
        if not glossary_terms:
            glossary_terms = list(DEFAULT_GLOSSARY_ENTRIES)
        
        # Get intensity metrics for chart generation
        intensity_metrics = get_intensity_metrics(job_id)
        
        # Generate ALL chart assets using the comprehensive function
        logger.debug("Generating report assets for job %s", job_id)
        assets = generate_report_assets(
            job_data=job_data,
            scope_totals=scope_totals,
            categories=categories,
            activity_totals=activity_totals,
            intensity_metrics=intensity_metrics,
            target_data=target_data
        )
        logger.debug("Generated report asset keys: %s", list(assets.keys()))
        
        # Extract chart images from assets for template compatibility
        scope_chart_base64 = assets.get('scope_breakdown', '')
        activity_chart_base64 = assets.get('activity_breakdown', '')
        logger.debug("scope_chart_base64 length: %s", len(scope_chart_base64))
        logger.debug("activity_chart_base64 length: %s", len(activity_chart_base64))
        
        # Setup Jinja2 template environment
        template_dir = os.path.join(os.path.dirname(__file__), 'templates')
        env = Environment(loader=FileSystemLoader(template_dir))
        template = _resolve_template_for_render(env, selected_template)
        
        # Get sites for Reporting Sites table
        report_sites = get_job_sites(job_id)
        
        # Render HTML with data
        html_content = template.render(
            job_data=job_data,
            scope_totals=scope_totals,
            benchmark_totals=benchmark_totals,
            categories=categories,
            activity_groups=activity_groups,
            activity_totals=activity_totals,
            activity_details=activity_details,
            activity_group_order=ACTIVITY_GROUP_ORDER,
            activity_group_colors=ACTIVITY_GROUP_COLORS,
            activity_labels=activity_labels,
            activity_values=activity_values,
            scope_chart_base64=scope_chart_base64,
            activity_chart_base64=activity_chart_base64,
            assets=assets,
            generation_date=generation_date,
            template_variables=template_variables,
            report_metadata=report_metadata,
            job_actions=job_actions,
            nzi_logo_src=_get_nzi_logo_src(),
            client_logo_src=_get_client_logo_src(job_data.get("logo_url")),
            render_values=render_values,
            render_template=render_meta,
            target_data=target_data,
            intensity_metrics=intensity_metrics,
            employee_number=report_metadata.get('employee_number'),
            report_sites=report_sites,
            glossary_terms=glossary_terms,
            site_breakdowns=site_breakdowns,
            is_yoy_template=is_yoy_template,
            scope_comparison=scope_comparison,
            benchmark_activity_totals=benchmark_activity_totals,
            activity_comparison=activity_comparison,
            benchmark_site_breakdowns=benchmark_site_breakdowns,
            site_overall_comparison=site_overall_comparison,
        )

        html_content = _apply_braced_placeholder_substitution(html_content, render_values)
        
        # Return HTML
        return Response(
            content=html_content,
            media_type="text/html",
            headers={
                "Content-Disposition": f"inline; filename=job-{job_id}-emissions-report.html",
                "X-Rendered-Template-Key": str((selected_template or {}).get("template_key") or ""),
                "X-Rendered-Template-Name": str((selected_template or {}).get("template_name") or ""),
                "X-Rendered-Version-Id": str((selected_template or {}).get("version_id") or ""),
                "X-Rendered-Template-Source": template_source,
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate HTML report: {str(e)}")


# Shared implementation used by api.job_report_pdf_routes.
def _generate_professional_pdf_impl(
    job_id: int,
    request: Request,
    save_version: bool = Query(False),
    report_version_status: str = Query("review"),
    report_version_label: str | None = Query(None),
    report_version_notes: str | None = Query(None),
    _user: dict[str, str] = Depends(_current_user),
):
    """
    Generate a professional PDF report using DocRaptor service.
    This produces high-quality, audit-ready PDFs from HTML templates.
    """
    try:
        # Fetch data
        job_data = get_job_data(job_id)
        if not job_data:
            raise HTTPException(status_code=404, detail="Job not found")

        if not save_version:
            with get_conn() as con:
                frozen_snapshot = _load_latest_final_report_version_snapshot(con, int(job_id))
                if frozen_snapshot:
                    version_row, snapshot_payload = frozen_snapshot
                    file_payload = None
                    file_id = version_row.get("file_id")
                    if file_id is not None:
                        file_payload = _load_report_version_file_payload(con, int(job_id), int(file_id))
                    if file_payload:
                        headers = {
                            "Content-Disposition": f"inline; filename={file_payload['file_name']}",
                            "X-Report-Version-Saved": "false",
                            "X-Report-Snapshot-Hash": str(version_row.get("data_hash") or ""),
                            "X-Report-Version-Id": str(version_row.get("report_version_id") or ""),
                            "X-Report-Version-Number": str(version_row.get("version_number") or ""),
                            "X-Report-Version-Label": str(version_row.get("version_label") or ""),
                            "X-Report-Version-Status": str(version_row.get("status") or ""),
                            "X-Report-File-Id": str(version_row.get("file_id") or ""),
                        }
                        return Response(
                            content=file_payload["content"],
                            media_type=str(file_payload.get("media_type") or "application/octet-stream"),
                            headers=headers,
                        )
                    html_content = _render_report_snapshot_html(snapshot_payload)
                    pdf_bytes = _render_html_to_pdf_bytes(html_content)
                    headers = {
                        "Content-Disposition": f"inline; filename=job-{job_id}-emissions-report.pdf",
                        "X-Report-Version-Saved": "false",
                        "X-Report-Snapshot-Hash": str(version_row.get("data_hash") or ""),
                        "X-Report-Version-Id": str(version_row.get("report_version_id") or ""),
                        "X-Report-Version-Number": str(version_row.get("version_number") or ""),
                        "X-Report-Version-Label": str(version_row.get("version_label") or ""),
                        "X-Report-Version-Status": str(version_row.get("status") or ""),
                        "X-Report-File-Id": str(version_row.get("file_id") or ""),
                    }
                    return Response(
                        content=pdf_bytes,
                        media_type="application/pdf",
                        headers=headers,
                    )
        
        scope_totals = get_scope_totals(job_id)
        categories = get_emissions_by_category(job_id)
        benchmark_totals = get_benchmark_emissions(job_id, job_data.get('benchmark_year'))
        site_breakdowns = get_site_emissions_breakdowns(job_id)
        
        activity_groups, activity_totals, activity_details = _build_activity_grouping(categories)
        benchmark_job_id = _get_benchmark_job_id(job_id, job_data.get("benchmark_year"))
        benchmark_categories = get_emissions_by_category(benchmark_job_id) if benchmark_job_id else []
        _bg, benchmark_activity_totals, _bd = _build_activity_grouping(benchmark_categories) if benchmark_categories else ({}, {}, {})
        benchmark_site_breakdowns = get_site_emissions_breakdowns(benchmark_job_id) if benchmark_job_id else {
            "show_site_tables": False,
            "site_count": 0,
            "overall": [],
            "scope": [],
            "activity": [],
            "appendix_rows": [],
        }
        period_from = job_data.get("period_start")
        period_to = job_data.get("period_end")
        is_yoy_template = False
        scope_comparison = _build_scope_comparison(scope_totals, benchmark_totals)
        activity_comparison = _build_activity_comparison(activity_totals, benchmark_activity_totals)
        site_overall_comparison = _build_site_overall_comparison(site_breakdowns, benchmark_site_breakdowns)

        # Prepare data for charts
        activity_labels = [g for g in ACTIVITY_GROUP_ORDER if float(activity_totals.get(g, 0) or 0) > 0]
        if not activity_labels:
            activity_labels = ACTIVITY_GROUP_ORDER
        activity_values = [float(activity_totals.get(g, 0) or 0) for g in activity_labels]

        selected_template = _get_job_assigned_template_selection(int(job_id))
        template_variables = {}
        render_meta = None
        if selected_template and selected_template.get("version_id") is not None:
            template_variables = _get_template_variable_values_for_render(
                int(job_id),
                int(selected_template["template_id"]),
                int(selected_template["version_id"]),
            )
            render_meta = {
                "template_id": selected_template.get("template_id"),
                "template_key": selected_template.get("template_key"),
                "template_name": selected_template.get("template_name"),
                "template_type": selected_template.get("template_type"),
                "report_type": selected_template.get("report_type"),
                "version_id": selected_template.get("version_id"),
                "version_number": selected_template.get("version_number"),
                "version_label": selected_template.get("version_label"),
                "version_status": selected_template.get("version_status"),
                "source": "assignment",
            }

        generation_date = datetime.now().strftime('%d %B %Y')
        report_metadata = _get_job_report_metadata(
            int(job_id),
            updated_by=_user.get("email", "unknown"),
        )
        job_actions = get_job_report_actions_payload(int(job_id))
        render_values = _build_report_render_values(
            job_data=job_data,
            scope_totals=scope_totals,
            template_variables=template_variables,
            report_metadata=report_metadata,
            generation_date=generation_date,
        )
        

        target_data = get_job_target_data(job_id)
        glossary_terms = _ensure_glossary_cards_and_fetch(job_id, job_data, target_data.get('glossary_terms'))
        if not glossary_terms:
            glossary_terms = list(DEFAULT_GLOSSARY_ENTRIES)

        # Generate chart images for PDF
        # Scope chart
        scope_chart_data = {
            'Scope 1': scope_totals['Scope 1'],
            'Scope 2': scope_totals['Scope 2'],
            'Scope 3': scope_totals['Scope 3']
        }
        scope_colors = {
            'Scope 1': '#4472C4',
            'Scope 2': '#ED7D31',
            'Scope 3': '#70AD47'
        }
        scope_chart_base64 = create_donut_chart(
            scope_chart_data,
            scope_colors,
            f'Emissions by Scope (tCOâ‚‚e)',
            scope_totals['Total']
        )
        
        # Activity chart - using new professional donut style
        activity_chart_base64 = create_activity_donut(
            activity_totals,
            ACTIVITY_GROUP_COLORS,
            title="Emissions by Activity",
            center_label="tCO₂e",
            period_from=period_from,
            period_to=period_to
        )
        
        # Setup Jinja2 template environment
        template_dir = os.path.join(os.path.dirname(__file__), 'templates')
        env = Environment(loader=FileSystemLoader(template_dir))
        template = _resolve_template_for_render(env, selected_template)
        
        # Render HTML with data
        html_content = template.render(
            job_data=job_data,
            scope_totals=scope_totals,
            benchmark_totals=benchmark_totals,
            categories=categories,
            activity_groups=activity_groups,
            activity_totals=activity_totals,
            activity_details=activity_details,
            activity_group_order=ACTIVITY_GROUP_ORDER,
            activity_group_colors=ACTIVITY_GROUP_COLORS,
            activity_labels=activity_labels,
            activity_values=activity_values,
            scope_chart_base64=scope_chart_base64,
            activity_chart_base64=activity_chart_base64,
            generation_date=generation_date,
            template_variables=template_variables,
            report_metadata=report_metadata,
            job_actions=job_actions,
            nzi_logo_src=_get_nzi_logo_src(),
            client_logo_src=_get_client_logo_src(job_data.get("logo_url")),
            render_values=render_values,
            render_template=render_meta,
            glossary_terms=glossary_terms,
            site_breakdowns=site_breakdowns,
            is_yoy_template=is_yoy_template,
            scope_comparison=scope_comparison,
            benchmark_activity_totals=benchmark_activity_totals,
            activity_comparison=activity_comparison,
            benchmark_site_breakdowns=benchmark_site_breakdowns,
            site_overall_comparison=site_overall_comparison,
        )

        html_content = _apply_braced_placeholder_substitution(html_content, render_values)

        report_snapshot_payload, report_snapshot_hash = _build_report_version_snapshot(
            job_data=job_data,
            selected_template=render_meta,
            template_variables=template_variables,
            report_metadata=report_metadata,
            job_actions=job_actions,
            scope_totals=scope_totals,
            benchmark_totals=benchmark_totals,
            categories=categories,
            activity_totals=activity_totals,
            site_breakdowns=site_breakdowns,
            benchmark_site_breakdowns=benchmark_site_breakdowns,
            render_values=render_values,
            target_data=target_data,
            scope_comparison=scope_comparison,
            activity_comparison=activity_comparison,
            site_overall_comparison=site_overall_comparison,
            assets=assets,
            generation_date=generation_date,
        )
        report_snapshot_json = json.dumps(report_snapshot_payload, ensure_ascii=False, sort_keys=True, default=str)
        normalized_version_status = _normalize_report_version_status(report_version_status)

        # Import DocRaptor inside function to avoid startup overhead
        import docraptor
        
        # Initialize DocRaptor client
        doc_api = docraptor.DocApi()
        doc_api.api_client.configuration.username = DOCRAPTOR_API_KEY
        
        # Get base URL from environment or use default
        public_base_url = os.getenv('PUBLIC_BASE_URL', 'http://localhost:8001')
        
        # Create PDF using DocRaptor
        try:
            pdf_bytes = doc_api.create_doc(
                {
                    "test": _DOCRAPTOR_TEST_MODE,
                    "document_content": html_content,
                    "document_type": "pdf",
                    "name": f"job-{job_id}-emissions-report.pdf",
                    "prince_options": {
                        "media": "print",
                        "baseurl": public_base_url,
                        "javascript": False,
                    },
                }
            )

            saved_report_version: dict[str, Any] | None = None
            if save_version:
                with get_conn(autocommit=False) as con:
                    _ensure_job_files_table(con)
                    _ensure_report_versions_schema(con)
                    version_row = con.execute(
                        "SELECT COALESCE(MAX(version_number), 0) + 1 FROM job_report_versions WHERE job_id = %s",
                        [int(job_id)],
                    ).fetchone()
                    next_version_number = int(version_row[0]) if version_row and version_row[0] is not None else 1
                    saved_report_version = _store_report_version_artifact(
                        con=con,
                        job_data=job_data,
                        version_number=next_version_number,
                        pdf_bytes=pdf_bytes,
                        generated_by=_user.get("email", "unknown"),
                        status=normalized_version_status,
                        template_id=int(selected_template["template_id"]) if selected_template and selected_template.get("template_id") is not None else None,
                        template_version_id=int(selected_template["version_id"]) if selected_template and selected_template.get("version_id") is not None else None,
                        data_hash=report_snapshot_hash,
                        snapshot_json=report_snapshot_json,
                        version_label=report_version_label,
                        notes=report_version_notes,
                    )
                    record_audit_event(
                        con,
                        request=request,
                        actor=_user,
                        action="create",
                        entity_type="job_report_version",
                        entity_id=saved_report_version.get("report_version_id"),
                        client_id=int(job_data.get("client_db_id") or 0) if job_data.get("client_db_id") is not None else None,
                        job_id=int(job_id),
                        after={
                            "report_version_id": saved_report_version.get("report_version_id"),
                            "version_number": saved_report_version.get("version_number"),
                            "version_label": saved_report_version.get("version_label"),
                            "status": saved_report_version.get("status"),
                            "file_id": saved_report_version.get("file_id"),
                            "data_hash": report_snapshot_hash,
                        },
                        metadata={
                            "report_snapshot_hash": report_snapshot_hash,
                            "report_version_status": normalized_version_status,
                            "save_version": True,
                        },
                    )

            headers = {
                "Content-Disposition": f"inline; filename=job-{job_id}-emissions-report.pdf",
                "X-Report-Version-Saved": "true" if saved_report_version else "false",
                "X-Report-Snapshot-Hash": report_snapshot_hash,
            }
            if saved_report_version:
                headers.update(
                    {
                        "X-Report-Version-Id": str(saved_report_version.get("report_version_id") or ""),
                        "X-Report-Version-Number": str(saved_report_version.get("version_number") or ""),
                        "X-Report-Version-Label": str(saved_report_version.get("version_label") or ""),
                        "X-Report-Version-Status": str(saved_report_version.get("status") or ""),
                        "X-Report-File-Id": str(saved_report_version.get("file_id") or ""),
                    }
                )

            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers=headers,
            )

        except docraptor.rest.ApiException as error:
            raise HTTPException(
                status_code=500,
                detail=f"DocRaptor API error: {getattr(error, 'status', 'unknown')} - {getattr(error, 'reason', 'unknown')}",
            )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate professional PDF: {str(e)}")


def generate_job_report_docx(
    job_id: int,
    payload: GenerateReportPayload | None = None,
    _user: dict[str, str] = Depends(_current_user),
):
    """Generate a DOCX report export for a job using the assigned template/version."""
    # Import python-docx inside function to avoid startup overhead
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="DOCX export dependency is not installed. Add python-docx to requirements.",
        )

    try:
        job_data = get_job_data(job_id)
        if not job_data:
            raise HTTPException(status_code=404, detail="Job not found")

        frozen_snapshot = None
        with get_conn() as con:
            frozen_snapshot = _load_latest_final_report_version_snapshot(con, int(job_id))
        if frozen_snapshot:
            version_row, snapshot_payload = frozen_snapshot
            job_data = snapshot_payload.get("job_data") or job_data
            scope_totals = snapshot_payload.get("scope_totals") or {"Scope 1": 0.0, "Scope 2": 0.0, "Scope 3": 0.0, "Total": 0.0}
            categories = snapshot_payload.get("categories") or []
            _activity_groups, activity_totals, activity_details = _build_activity_grouping(categories)
            template_data = snapshot_payload.get("template") or {}
            selected_template = template_data if isinstance(template_data, dict) else {}
            template_variables = snapshot_payload.get("template_variables") or {}
            report_metadata = snapshot_payload.get("report_metadata") or {}
            job_actions = snapshot_payload.get("job_actions") or {}
            generation_date = snapshot_payload.get("generation_date") or datetime.now().strftime('%d %B %Y')
            render_values = snapshot_payload.get("render_values") or {}
            template_var_meta = _get_template_variable_metadata(int(selected_template.get("template_id") or 0)) if selected_template.get("template_id") is not None else []

            template_name = str(selected_template.get("template_name") or "")
            template_type = str(selected_template.get("template_type") or "")
            report_type = str(selected_template.get("report_type") or "")
            version_number = selected_template.get("version_number")
            template_hint = f"{template_name} {template_type} {report_type}".lower()
            is_annual = ("annual" in template_hint) or ("carbon_report" in template_hint)
            is_crp = ("crp" in template_hint) or ("carbon_reduction" in template_hint)

            doc = DocxDocument()
            _apply_docx_brand_styles(doc)

            # Cover page
            _docx_add_cover_page(
                doc,
                job_data=job_data,
                generation_date=generation_date,
                template_name=template_name,
                report_metadata=report_metadata,
            )

            if is_crp:
                _docx_add_section_heading(doc, "Carbon Reduction Plan Summary", level=1)
                crp_fields = [
                    ("Executive Summary", template_variables.get("executive_summary")),
                    ("Commitment Statement", template_variables.get("commitment_statement")),
                    ("Reduction Targets", template_variables.get("reduction_targets")),
                    ("Reduction Projects", template_variables.get("reduction_projects")),
                ]
                for label, raw in crp_fields:
                    value = _format_template_value(raw, "textarea")
                    if not value:
                        continue
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(value)
            elif is_annual:
                _docx_add_section_heading(doc, "Annual Report Narrative", level=1)
                annual_fields = [
                    ("Introduction", template_variables.get("introduction")),
                    ("Methodology", template_variables.get("methodology")),
                    ("Key Findings", template_variables.get("key_findings")),
                    ("Recommendations", template_variables.get("recommendations")),
                    ("Conclusion", template_variables.get("conclusion")),
                ]
                for label, raw in annual_fields:
                    value = _format_template_value(raw, "textarea")
                    if not value:
                        continue
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(value)

            # â”€â”€ Section: Narrative â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
            if is_crp:
                doc.add_heading("Carbon Reduction Plan", level=1)
                crp_fields = [
                    ("Executive Summary", template_variables.get("executive_summary")),
                    ("Commitment Statement", template_variables.get("commitment_statement")),
                    ("Reduction Targets", template_variables.get("reduction_targets")),
                    ("Reduction Projects", template_variables.get("reduction_projects")),
                ]
                for label, raw in crp_fields:
                    value = _format_template_value(raw, "textarea")
                    if not value:
                        continue
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(value)
            elif is_annual:
                doc.add_heading("Report Narrative", level=1)
                annual_fields = [
                    ("Introduction", template_variables.get("introduction")),
                    ("Methodology", template_variables.get("methodology")),
                    ("Key Findings", template_variables.get("key_findings")),
                    ("Recommendations", template_variables.get("recommendations")),
                    ("Conclusion", template_variables.get("conclusion")),
                ]
                for label, raw in annual_fields:
                    value = _format_template_value(raw, "textarea")
                    if not value:
                        continue
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(value)

            # â”€â”€ Section: Charts â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
            snapshot_assets = snapshot_payload.get("assets") or {}
            scope_chart = (snapshot_assets.get("total_emissions") or
                           snapshot_assets.get("scope_breakdown") or
                           snapshot_payload.get("scope_chart_base64") or "")
            activity_chart = (snapshot_assets.get("activity_breakdown") or
                              snapshot_payload.get("activity_chart_base64") or "")
            if scope_chart or activity_chart:
                doc.add_heading("Emissions Charts", level=1)
                _docx_embed_chart(doc, scope_chart, "Scope Profile (tCOâ‚‚e)")
                _docx_embed_chart(doc, activity_chart, "Activity Profile (tCOâ‚‚e)")

            # â”€â”€ Section: Emissions Summary â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
            doc.add_heading("Emissions Summary", level=1)
            summary_table = doc.add_table(rows=1, cols=3)
            summary_table.style = "Table Grid"
            for col, hdr in enumerate(["Scope", "tCOâ‚‚e", "% of Total"]):
                summary_table.cell(0, col).text = hdr
            _docx_shade_table_header(summary_table)
            s1 = scope_totals.get('Scope 1', 0.0)
            s2 = scope_totals.get('Scope 2', 0.0)
            s3 = scope_totals.get('Scope 3', 0.0)
            total = scope_totals.get('Total', 0.0)
            for scope_label, val in [("Scope 1", s1), ("Scope 2", s2), ("Scope 3", s3), ("Total", total)]:
                rc = summary_table.add_row().cells
                rc[0].text = scope_label
                rc[1].text = f"{val:,.2f}"
                rc[2].text = f"{(val / total * 100):.1f}%" if total > 0 else "0.0%"

            # â”€â”€ Section: Activity Group Totals â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
            doc.add_heading("Activity Group Totals", level=1)
            group_table = doc.add_table(rows=1, cols=2)
            group_table.style = "Table Grid"
            group_table.cell(0, 0).text = "Activity Group"
            group_table.cell(0, 1).text = "tCOâ‚‚e"
            _docx_shade_table_header(group_table)
            for group in ACTIVITY_GROUP_ORDER:
                rc = group_table.add_row().cells
                rc[0].text = group
                rc[1].text = f"{activity_totals.get(group, 0.0):,.2f}"

            # â”€â”€ Section: Planned Actions â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
            if job_actions.get("items"):
                doc.add_heading("Planned Actions", level=1)
                actions_table = doc.add_table(rows=1, cols=4)
                actions_table.style = "Table Grid"
                for col, hdr in enumerate(["Term", "Action", "Category", "Description"]):
                    actions_table.cell(0, col).text = hdr
                _docx_shade_table_header(actions_table)
                for item in job_actions.get("items", []):
                    row_cells = actions_table.add_row().cells
                    row_cells[0].text = str(item.get("action_term_label") or item.get("action_term") or "")
                    row_cells[1].text = str(item.get("action_name") or "")
                    row_cells[2].text = str(item.get("action_category") or "")
                    row_cells[3].text = str(item.get("description") or "")

            # â”€â”€ Section: Top Activity Rows â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
            if activity_details:
                doc.add_heading("Top Emission Drivers", level=1)
                details_table = doc.add_table(rows=1, cols=5)
                details_table.style = "Table Grid"
                for col, hdr in enumerate(["Activity Group", "Emission Type", "Scope", "Confidence", "tCOâ‚‚e"]):
                    details_table.cell(0, col).text = hdr
                _docx_shade_table_header(details_table)
                for row in activity_details[:50]:
                    cells = details_table.add_row().cells
                    cells[0].text = str(row.get("activity_group") or "")
                    cells[1].text = str(row.get("emission_type") or row.get("report_label") or "")
                    cells[2].text = str(row.get("scope") or "")
                    cells[3].text = str(row.get("data_confidence") or "M")
                    cells[4].text = f"{float(row.get('emissions') or 0.0):,.2f}"

            # â”€â”€ Section: Additional template variable sections â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
            _docx_add_variable_sections(doc, template_variables, template_var_meta)

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            file_name = _safe_filename(
                f"job-{job_data.get('job_number') or job_id}-emissions-report",
                "docx",
            )

            return Response(
                content=output.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f'attachment; filename="{file_name}"',
                    "X-Report-Version-Saved": "false",
                    "X-Report-Snapshot-Hash": str(version_row.get("data_hash") or ""),
                    "X-Report-Version-Id": str(version_row.get("report_version_id") or ""),
                    "X-Report-Version-Number": str(version_row.get("version_number") or ""),
                    "X-Report-Version-Label": str(version_row.get("version_label") or ""),
                    "X-Report-Version-Status": str(version_row.get("status") or ""),
                    "X-Report-File-Id": str(version_row.get("file_id") or ""),
                },
            )

        scope_totals = get_scope_totals(job_id)
        categories = get_emissions_by_category(job_id)
        _activity_groups, activity_totals, activity_details = _build_activity_grouping(categories)

        selected_template = _resolve_selected_template(int(job_id), payload)

        template_variables = _get_template_variable_values_for_render(
            int(job_id),
            int(selected_template["template_id"]),
            int(selected_template["version_id"]),
        )
        report_metadata = _get_job_report_metadata(
            int(job_id),
            updated_by=_user.get("email", "unknown"),
        )
        job_actions = get_job_report_actions_payload(int(job_id))
        generation_date = datetime.now().strftime('%d %B %Y')
        render_values = _build_report_render_values(
            job_data=job_data,
            scope_totals=scope_totals,
            template_variables=template_variables,
            report_metadata=report_metadata,
            generation_date=generation_date,
        )
        template_var_meta = _get_template_variable_metadata(int(selected_template["template_id"]))

        template_name = str(selected_template.get("template_name") or "")
        template_type = str(selected_template.get("template_type") or "")
        report_type = str(selected_template.get("report_type") or "")
        version_number = selected_template.get("version_number")

        template_hint = f"{template_name} {template_type} {report_type}".lower()
        is_annual = ("annual" in template_hint) or ("carbon_report" in template_hint)
        is_crp = ("crp" in template_hint) or ("carbon_reduction" in template_hint)

        doc = DocxDocument()

        heading = doc.add_heading(
            str(
                render_values.get("report_title")
                or template_variables.get("report_title")
                or f"Emissions Report â€“ {job_data.get('client_name') or 'Client'}"
            ),
            level=0,
        )
        if WD_PARAGRAPH_ALIGNMENT is not None:
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle = doc.add_paragraph(
            f"Job {job_data.get('job_number') or job_id} â€¢ Reporting year {job_data.get('reporting_year') or 'N/A'}"
        )
        if WD_PARAGRAPH_ALIGNMENT is not None:
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(
            f"Generated: {generation_date}"
        )

        meta_p = doc.add_paragraph(
            f"Template: {template_name or selected_template.get('template_key') or selected_template.get('template_id')}"
        )
        if version_number is not None:
            meta_p.add_run(f" (v{version_number})")

        if is_crp:
            _docx_add_section_heading(doc, "Carbon Reduction Plan Summary", level=1)
            crp_fields = [
                ("Executive Summary", template_variables.get("executive_summary")),
                ("Commitment Statement", template_variables.get("commitment_statement")),
                ("Reduction Targets", template_variables.get("reduction_targets")),
                ("Reduction Projects", template_variables.get("reduction_projects")),
            ]
            for label, raw in crp_fields:
                value = _format_template_value(raw, "textarea")
                if not value:
                    continue
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(value)
        elif is_annual:
            _docx_add_section_heading(doc, "Annual Report Narrative", level=1)
            annual_fields = [
                ("Introduction", template_variables.get("introduction")),
                ("Methodology", template_variables.get("methodology")),
                ("Key Findings", template_variables.get("key_findings")),
                ("Recommendations", template_variables.get("recommendations")),
                ("Conclusion", template_variables.get("conclusion")),
            ]
            for label, raw in annual_fields:
                value = _format_template_value(raw, "textarea")
                if not value:
                    continue
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(value)

        if job_actions.get("items"):
            doc.add_heading("Planned Actions", level=1)
            actions_table = doc.add_table(rows=1, cols=4)
            actions_table.style = "Table Grid"
            action_headers = ["Term", "Action", "Category", "Description"]
            for col, header in enumerate(action_headers):
                actions_table.cell(0, col).text = header

            for item in job_actions.get("items", []):
                row_cells = actions_table.add_row().cells
                row_cells[0].text = str(item.get("action_term_label") or item.get("action_term") or "")
                row_cells[1].text = str(item.get("action_name") or "")
                row_cells[2].text = str(item.get("action_category") or "")
                row_cells[3].text = str(item.get("description") or "")

        doc.add_heading("Summary", level=1)
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = "Table Grid"
        summary_rows = [
            ("Scope 1", f"{scope_totals.get('Scope 1', 0.0):,.2f} tCOâ‚‚e"),
            ("Scope 2", f"{scope_totals.get('Scope 2', 0.0):,.2f} tCOâ‚‚e"),
            ("Scope 3", f"{scope_totals.get('Scope 3', 0.0):,.2f} tCOâ‚‚e"),
            ("Total", f"{scope_totals.get('Total', 0.0):,.2f} tCOâ‚‚e"),
        ]
        for idx, (label, value) in enumerate(summary_rows):
            summary_table.cell(idx, 0).text = label
            summary_table.cell(idx, 1).text = value

        doc.add_heading("Activity Group Totals", level=1)
        group_table = doc.add_table(rows=max(1, len(ACTIVITY_GROUP_ORDER)), cols=2)
        group_table.style = "Table Grid"
        for idx, group in enumerate(ACTIVITY_GROUP_ORDER):
            group_table.cell(idx, 0).text = group
            group_table.cell(idx, 1).text = f"{activity_totals.get(group, 0.0):,.2f} tCOâ‚‚e"

        _docx_add_variable_sections(doc, template_variables, template_var_meta)

        metadata_rows = [
            (key, report_metadata.get(key))
            for key in report_metadata.keys()
            if report_metadata.get(key) not in (None, "")
        ]
        if metadata_rows:
            doc.add_heading("Report Metadata", level=1)
            meta_table = doc.add_table(rows=1, cols=2)
            meta_table.style = "Table Grid"
            meta_table.cell(0, 0).text = "Field"
            meta_table.cell(0, 1).text = "Value"
            for key, value in metadata_rows:
                row_cells = meta_table.add_row().cells
                row_cells[0].text = str(key).replace("_", " ").title()
                row_cells[1].text = _stringify_render_value(value)

        if activity_details:
            doc.add_heading("Top Activity Rows", level=1)
            details_table = doc.add_table(rows=1, cols=4)
            details_table.style = "Table Grid"
            headers = ["Group", "Emission Type", "Scope", "tCOâ‚‚e"]
            for col, header in enumerate(headers):
                details_table.cell(0, col).text = header

            for row in activity_details[:50]:
                cells = details_table.add_row().cells
                cells[0].text = str(row.get("activity_group") or "")
                cells[1].text = str(row.get("emission_type") or "")
                cells[2].text = str(row.get("scope") or "")
                cells[3].text = f"{float(row.get('emissions') or 0.0):,.2f}"

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        file_name = _safe_filename(
            f"job-{job_data.get('job_number') or job_id}-emissions-report",
            "docx",
        )

        return Response(
            content=output.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{file_name}"',
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate DOCX report: {str(e)}")
