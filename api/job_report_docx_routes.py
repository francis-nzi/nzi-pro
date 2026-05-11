from __future__ import annotations

import io
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response

from api.auth import _current_user
from api.job_report_routes import (
    GenerateReportPayload,
    _apply_docx_brand_styles,
    _build_activity_grouping,
    _build_report_render_values,
    _docx_add_cover_page,
    _docx_add_section_heading,
    _docx_add_variable_sections,
    _docx_embed_chart,
    _docx_shade_table_header,
    _format_template_value,
    _get_job_report_metadata,
    _get_template_variable_metadata,
    _get_template_variable_values_for_render,
    _load_latest_final_report_version_snapshot,
    _normalize_report_version_status,
    _resolve_selected_template,
    _safe_filename,
    _stringify_render_value,
    get_emissions_by_category,
    get_job_data,
    get_job_report_actions_payload,
    get_scope_totals,
)
from core.database import get_conn

router = APIRouter()


@router.post("/jobs/{job_id}/generate-report-docx")
def generate_job_report_docx(
    job_id: int,
    payload: GenerateReportPayload | None = None,
    _user: dict[str, str] = Depends(_current_user),
):
    """Generate a DOCX report export for a job using the assigned template/version."""
    try:
        from docx import Document as DocxDocument
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="DOCX export dependency is not installed. Add python-docx to requirements.",
        )

    try:
        job_data = get_job_data(job_id)
        if not job_data:
            raise HTTPException(status_code=404, detail="Job not found")

        frozen_snapshot = None
        with get_conn() as con:
            frozen_snapshot = _load_latest_final_report_version_snapshot(con, int(job_id))
        if frozen_snapshot:
            version_row, snapshot_payload = frozen_snapshot
            job_data = snapshot_payload.get("job_data") or job_data
            scope_totals = snapshot_payload.get("scope_totals") or {"Scope 1": 0.0, "Scope 2": 0.0, "Scope 3": 0.0, "Total": 0.0}
            categories = snapshot_payload.get("categories") or []
            _activity_groups, activity_totals, activity_details = _build_activity_grouping(categories)
            template_data = snapshot_payload.get("template") or {}
            selected_template = template_data if isinstance(template_data, dict) else {}
            template_variables = snapshot_payload.get("template_variables") or {}
            report_metadata = snapshot_payload.get("report_metadata") or {}
            job_actions = snapshot_payload.get("job_actions") or {}
            generation_date = snapshot_payload.get("generation_date") or datetime.now().strftime('%d %B %Y')
            render_values = snapshot_payload.get("render_values") or {}
            template_var_meta = _get_template_variable_metadata(int(selected_template.get("template_id") or 0)) if selected_template.get("template_id") is not None else []

            template_name = str(selected_template.get("template_name") or "")
            template_type = str(selected_template.get("template_type") or "")
            report_type = str(selected_template.get("report_type") or "")
            version_number = selected_template.get("version_number")
            template_hint = f"{template_name} {template_type} {report_type}".lower()
            is_annual = ("annual" in template_hint) or ("carbon_report" in template_hint)
            is_crp = ("crp" in template_hint) or ("carbon_reduction" in template_hint)

            doc = DocxDocument()
            _apply_docx_brand_styles(doc)

            _docx_add_cover_page(
                doc,
                job_data=job_data,
                generation_date=generation_date,
                template_name=template_name,
                report_metadata=report_metadata,
            )

            if is_crp:
                _docx_add_section_heading(doc, "Carbon Reduction Plan Summary", level=1)
                crp_fields = [
                    ("Executive Summary", template_variables.get("executive_summary")),
                    ("Commitment Statement", template_variables.get("commitment_statement")),
                    ("Reduction Targets", template_variables.get("reduction_targets")),
                    ("Reduction Projects", template_variables.get("reduction_projects")),
                ]
                for label, raw in crp_fields:
                    value = _format_template_value(raw, "textarea")
                    if not value:
                        continue
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(value)
            elif is_annual:
                _docx_add_section_heading(doc, "Annual Report Narrative", level=1)
                annual_fields = [
                    ("Introduction", template_variables.get("introduction")),
                    ("Methodology", template_variables.get("methodology")),
                    ("Key Findings", template_variables.get("key_findings")),
                    ("Recommendations", template_variables.get("recommendations")),
                    ("Conclusion", template_variables.get("conclusion")),
                ]
                for label, raw in annual_fields:
                    value = _format_template_value(raw, "textarea")
                    if not value:
                        continue
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(value)

            if is_crp:
                doc.add_heading("Carbon Reduction Plan", level=1)
                crp_fields = [
                    ("Executive Summary", template_variables.get("executive_summary")),
                    ("Commitment Statement", template_variables.get("commitment_statement")),
                    ("Reduction Targets", template_variables.get("reduction_targets")),
                    ("Reduction Projects", template_variables.get("reduction_projects")),
                ]
                for label, raw in crp_fields:
                    value = _format_template_value(raw, "textarea")
                    if not value:
                        continue
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(value)
            elif is_annual:
                doc.add_heading("Report Narrative", level=1)
                annual_fields = [
                    ("Introduction", template_variables.get("introduction")),
                    ("Methodology", template_variables.get("methodology")),
                    ("Key Findings", template_variables.get("key_findings")),
                    ("Recommendations", template_variables.get("recommendations")),
                    ("Conclusion", template_variables.get("conclusion")),
                ]
                for label, raw in annual_fields:
                    value = _format_template_value(raw, "textarea")
                    if not value:
                        continue
                    p = doc.add_paragraph()
                    p.add_run(f"{label}: ").bold = True
                    p.add_run(value)

            snapshot_assets = snapshot_payload.get("assets") or {}
            scope_chart = (
                snapshot_assets.get("total_emissions")
                or snapshot_assets.get("scope_breakdown")
                or snapshot_payload.get("scope_chart_base64")
                or ""
            )
            activity_chart = (
                snapshot_assets.get("activity_breakdown")
                or snapshot_payload.get("activity_chart_base64")
                or ""
            )
            if scope_chart or activity_chart:
                doc.add_heading("Emissions Charts", level=1)
                _docx_embed_chart(doc, scope_chart, "Scope Profile (tCOâ‚‚e)")
                _docx_embed_chart(doc, activity_chart, "Activity Profile (tCOâ‚‚e)")

            doc.add_heading("Emissions Summary", level=1)
            summary_table = doc.add_table(rows=1, cols=3)
            summary_table.style = "Table Grid"
            for col, hdr in enumerate(["Scope", "tCOâ‚‚e", "% of Total"]):
                summary_table.cell(0, col).text = hdr
            _docx_shade_table_header(summary_table)
            s1 = scope_totals.get('Scope 1', 0.0)
            s2 = scope_totals.get('Scope 2', 0.0)
            s3 = scope_totals.get('Scope 3', 0.0)
            total = scope_totals.get('Total', 0.0)
            for scope_label, val in [("Scope 1", s1), ("Scope 2", s2), ("Scope 3", s3), ("Total", total)]:
                rc = summary_table.add_row().cells
                rc[0].text = scope_label
                rc[1].text = f"{val:,.2f}"
                rc[2].text = f"{(val / total * 100):.1f}%" if total > 0 else "0.0%"

            doc.add_heading("Activity Group Totals", level=1)
            group_table = doc.add_table(rows=1, cols=2)
            group_table.style = "Table Grid"
            group_table.cell(0, 0).text = "Activity Group"
            group_table.cell(0, 1).text = "tCOâ‚‚e"
            _docx_shade_table_header(group_table)
            activity_groups, activity_totals, activity_details = _build_activity_grouping(categories)
            for group in ["Energy", "Business Travel", "Employee Commuting", "Purchased Goods & Services (PG&S)", "Other Emissions"]:
                rc = group_table.add_row().cells
                rc[0].text = group
                rc[1].text = f"{activity_totals.get(group, 0.0):,.2f}"

            if job_actions.get("items"):
                doc.add_heading("Planned Actions", level=1)
                actions_table = doc.add_table(rows=1, cols=4)
                actions_table.style = "Table Grid"
                for col, hdr in enumerate(["Term", "Action", "Category", "Description"]):
                    actions_table.cell(0, col).text = hdr
                _docx_shade_table_header(actions_table)
                for item in job_actions.get("items", []):
                    row_cells = actions_table.add_row().cells
                    row_cells[0].text = str(item.get("action_term_label") or item.get("action_term") or "")
                    row_cells[1].text = str(item.get("action_name") or "")
                    row_cells[2].text = str(item.get("action_category") or "")
                    row_cells[3].text = str(item.get("description") or "")

            if activity_details:
                doc.add_heading("Top Emission Drivers", level=1)
                details_table = doc.add_table(rows=1, cols=5)
                details_table.style = "Table Grid"
                for col, hdr in enumerate(["Activity Group", "Emission Type", "Scope", "Confidence", "tCOâ‚‚e"]):
                    details_table.cell(0, col).text = hdr
                _docx_shade_table_header(details_table)
                for row in activity_details[:50]:
                    cells = details_table.add_row().cells
                    cells[0].text = str(row.get("activity_group") or "")
                    cells[1].text = str(row.get("emission_type") or row.get("report_label") or "")
                    cells[2].text = str(row.get("scope") or "")
                    cells[3].text = str(row.get("data_confidence") or "M")
                    cells[4].text = f"{float(row.get('emissions') or 0.0):,.2f}"

            _docx_add_variable_sections(doc, template_variables, template_var_meta)

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            file_name = _safe_filename(
                f"job-{job_data.get('job_number') or job_id}-emissions-report",
                "docx",
            )

            return Response(
                content=output.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f'attachment; filename="{file_name}"',
                    "X-Report-Version-Saved": "false",
                    "X-Report-Snapshot-Hash": str(version_row.get("data_hash") or ""),
                    "X-Report-Version-Id": str(version_row.get("report_version_id") or ""),
                    "X-Report-Version-Number": str(version_row.get("version_number") or ""),
                    "X-Report-Version-Label": str(version_row.get("version_label") or ""),
                    "X-Report-Version-Status": str(version_row.get("status") or ""),
                    "X-Report-File-Id": str(version_row.get("file_id") or ""),
                },
            )

        scope_totals = get_scope_totals(job_id)
        categories = get_emissions_by_category(job_id)
        _activity_groups, activity_totals, activity_details = _build_activity_grouping(categories)

        selected_template = _resolve_selected_template(int(job_id), payload)

        template_variables = _get_template_variable_values_for_render(
            int(job_id),
            int(selected_template["template_id"]),
            int(selected_template["version_id"]),
        )
        report_metadata = _get_job_report_metadata(
            int(job_id),
            updated_by=_user.get("email", "unknown"),
        )
        job_actions = get_job_report_actions_payload(int(job_id))
        generation_date = datetime.now().strftime('%d %B %Y')
        render_values = _build_report_render_values(
            job_data=job_data,
            scope_totals=scope_totals,
            template_variables=template_variables,
            report_metadata=report_metadata,
            generation_date=generation_date,
        )
        template_var_meta = _get_template_variable_metadata(int(selected_template["template_id"]))

        template_name = str(selected_template.get("template_name") or "")
        template_type = str(selected_template.get("template_type") or "")
        report_type = str(selected_template.get("report_type") or "")
        version_number = selected_template.get("version_number")

        template_hint = f"{template_name} {template_type} {report_type}".lower()
        is_annual = ("annual" in template_hint) or ("carbon_report" in template_hint)
        is_crp = ("crp" in template_hint) or ("carbon_reduction" in template_hint)

        doc = DocxDocument()

        heading = doc.add_heading(
            str(
                render_values.get("report_title")
                or template_variables.get("report_title")
                or f"Emissions Report â€“ {job_data.get('client_name') or 'Client'}"
            ),
            level=0,
        )
        if WD_PARAGRAPH_ALIGNMENT is not None:
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle = doc.add_paragraph(
            f"Job {job_data.get('job_number') or job_id} â€¢ Reporting year {job_data.get('reporting_year') or 'N/A'}"
        )
        if WD_PARAGRAPH_ALIGNMENT is not None:
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(f"Generated: {generation_date}")

        meta_p = doc.add_paragraph(
            f"Template: {template_name or selected_template.get('template_key') or selected_template.get('template_id')}"
        )
        if version_number is not None:
            meta_p.add_run(f" (v{version_number})")

        if is_crp:
            _docx_add_section_heading(doc, "Carbon Reduction Plan Summary", level=1)
            crp_fields = [
                ("Executive Summary", template_variables.get("executive_summary")),
                ("Commitment Statement", template_variables.get("commitment_statement")),
                ("Reduction Targets", template_variables.get("reduction_targets")),
                ("Reduction Projects", template_variables.get("reduction_projects")),
            ]
            for label, raw in crp_fields:
                value = _format_template_value(raw, "textarea")
                if not value:
                    continue
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(value)
        elif is_annual:
            _docx_add_section_heading(doc, "Annual Report Narrative", level=1)
            annual_fields = [
                ("Introduction", template_variables.get("introduction")),
                ("Methodology", template_variables.get("methodology")),
                ("Key Findings", template_variables.get("key_findings")),
                ("Recommendations", template_variables.get("recommendations")),
                ("Conclusion", template_variables.get("conclusion")),
            ]
            for label, raw in annual_fields:
                value = _format_template_value(raw, "textarea")
                if not value:
                    continue
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(value)

        if job_actions.get("items"):
            doc.add_heading("Planned Actions", level=1)
            actions_table = doc.add_table(rows=1, cols=4)
            actions_table.style = "Table Grid"
            action_headers = ["Term", "Action", "Category", "Description"]
            for col, header in enumerate(action_headers):
                actions_table.cell(0, col).text = header

            for item in job_actions.get("items", []):
                row_cells = actions_table.add_row().cells
                row_cells[0].text = str(item.get("action_term_label") or item.get("action_term") or "")
                row_cells[1].text = str(item.get("action_name") or "")
                row_cells[2].text = str(item.get("action_category") or "")
                row_cells[3].text = str(item.get("description") or "")

        doc.add_heading("Summary", level=1)
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = "Table Grid"
        summary_rows = [
            ("Scope 1", f"{scope_totals.get('Scope 1', 0.0):,.2f} tCOâ‚‚e"),
            ("Scope 2", f"{scope_totals.get('Scope 2', 0.0):,.2f} tCOâ‚‚e"),
            ("Scope 3", f"{scope_totals.get('Scope 3', 0.0):,.2f} tCOâ‚‚e"),
            ("Total", f"{scope_totals.get('Total', 0.0):,.2f} tCOâ‚‚e"),
        ]
        for idx, (label, value) in enumerate(summary_rows):
            summary_table.cell(idx, 0).text = label
            summary_table.cell(idx, 1).text = value

        doc.add_heading("Activity Group Totals", level=1)
        group_table = doc.add_table(rows=max(1, len(["Energy", "Business Travel", "Employee Commuting", "Purchased Goods & Services (PG&S)", "Other Emissions"])), cols=2)
        group_table.style = "Table Grid"
        for idx, group in enumerate(["Energy", "Business Travel", "Employee Commuting", "Purchased Goods & Services (PG&S)", "Other Emissions"]):
            group_table.cell(idx, 0).text = group
            group_table.cell(idx, 1).text = f"{activity_totals.get(group, 0.0):,.2f} tCOâ‚‚e"

        _docx_add_variable_sections(doc, template_variables, template_var_meta)

        metadata_rows = [
            (key, report_metadata.get(key))
            for key in report_metadata.keys()
            if report_metadata.get(key) not in (None, "")
        ]
        if metadata_rows:
            doc.add_heading("Report Metadata", level=1)
            meta_table = doc.add_table(rows=1, cols=2)
            meta_table.style = "Table Grid"
            meta_table.cell(0, 0).text = "Field"
            meta_table.cell(0, 1).text = "Value"
            for key, value in metadata_rows:
                row_cells = meta_table.add_row().cells
                row_cells[0].text = str(key).replace("_", " ").title()
                row_cells[1].text = _stringify_render_value(value)

        if activity_details:
            doc.add_heading("Top Activity Rows", level=1)
            details_table = doc.add_table(rows=1, cols=4)
            details_table.style = "Table Grid"
            headers = ["Group", "Emission Type", "Scope", "tCOâ‚‚e"]
            for col, header in enumerate(headers):
                details_table.cell(0, col).text = header

            for row in activity_details[:50]:
                cells = details_table.add_row().cells
                cells[0].text = str(row.get("activity_group") or "")
                cells[1].text = str(row.get("emission_type") or "")
                cells[2].text = str(row.get("scope") or "")
                cells[3].text = f"{float(row.get('emissions') or 0.0):,.2f}"

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        file_name = _safe_filename(
            f"job-{job_data.get('job_number') or job_id}-emissions-report",
            "docx",
        )

        return Response(
            content=output.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{file_name}"',
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate DOCX report: {str(e)}")
